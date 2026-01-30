import sqlite3
import os
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import json

# ========== ФУНКЦИИ БАЗЫ ДАННЫХ ==========

def init_db():
    """Инициализация базы данных"""
    conn = sqlite3.connect("достижения.db")
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS достижения(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        название TEXT NOT NULL,
        дата TEXT NOT NULL,
        тип TEXT NOT NULL,
        уровень TEXT NOT NULL,
        описание TEXT
    )
    """)
    conn.commit()
    conn.close()


def save_to_db(name, date, typ, level, desc):
    """Сохранение записи в базу данных"""
    conn = sqlite3.connect("достижения.db")
    cur = conn.cursor()
    cur.execute("INSERT INTO достижения (название, дата, тип, уровень, описание) VALUES (?, ?, ?, ?, ?)",
                (name, date, typ, level, desc))
    conn.commit()
    conn.close()


def load_records():
    """Загрузка записей из базы данных"""
    conn = sqlite3.connect("достижения.db")
    cur = conn.cursor()
    cur.execute("SELECT дата, название, тип, уровень FROM достижения ORDER BY дата DESC")
    rows = cur.fetchall()
    conn.close()
    return rows


def load_records_with_desc():
    """Загрузка записей с описанием"""
    conn = sqlite3.connect("достижения.db")
    cur = conn.cursor()
    cur.execute("SELECT дата, название, тип, уровень, описание FROM достижения ORDER BY дата DESC")
    rows = cur.fetchall()
    conn.close()
    return rows


def delete_record(selected_index):
    """Удаление записи из базы данных"""
    if selected_index:
        conn = sqlite3.connect("достижения.db")
        cur = conn.cursor()
        # Получаем ID записи для удаления
        cur.execute("SELECT id FROM достижения ORDER BY дата DESC LIMIT 1 OFFSET ?", (selected_index,))
        record_id = cur.fetchone()
        if record_id:
            cur.execute("DELETE FROM достижения WHERE id = ?", (record_id[0],))
            conn.commit()
        conn.close()
        return True
    return False


# ========== ФУНКЦИИ ДЛЯ РАБОТЫ С ФАЙЛАМИ ==========

def load_types():
    """Загрузка типов достижений из JSON файла"""
    try:
        with open("types.json", "r", encoding="utf-8") as f:
            data = json.load(f)
            # Фильтруем пустые строки и приводим к правильному формату
            return [item.strip().title() for item in data if item and str(item).strip()]
    except Exception as e:
        print(f"Ошибка загрузки types.json: {e}")
        return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]


# ========== ФУНКЦИИ ГРАФИЧЕСКОГО ИНТЕРФЕЙСА ==========

def create_add_form(parent):
    """Создание формы для добавления достижений"""
    # Стилизация
    style = ttk.Style()
    style.configure("Custom.TLabel", font=("Arial", 10), padding=5)
    style.configure("Custom.TEntry", font=("Arial", 10), padding=5)

    # Основной контейнер с отступами
    main_frame = ttk.Frame(parent)
    main_frame.pack(padx=20, pady=20, fill="both", expand=True)

    # Поле "Название"
    tk.Label(main_frame, text="Название:", font=("Arial", 10, "bold"),
             bg="#f0f0f0", anchor="w").pack(fill="x", padx=5, pady=(0, 5))
    name_entry = tk.Entry(main_frame, width=60, font=("Arial", 10),
                          relief="solid", bd=1)
    name_entry.pack(padx=5, pady=(0, 15))

    # Поле "Дата"
    tk.Label(main_frame, text="Дата (ГГГГ-ММ-ДД):", font=("Arial", 10, "bold"),
             bg="#f0f0f0", anchor="w").pack(fill="x", padx=5, pady=(0, 5))
    date_entry = tk.Entry(main_frame, width=60, font=("Arial", 10),
                          relief="solid", bd=1)
    date_entry.pack(padx=5, pady=(0, 15))

    # Поле "Тип"
    tk.Label(main_frame, text="Тип:", font=("Arial", 10, "bold"),
             bg="#f0f0f0", anchor="w").pack(fill="x", padx=5, pady=(0, 5))
    types = load_types()
    type_combo = ttk.Combobox(main_frame, values=types, state="readonly",
                              font=("Arial", 10), width=58)
    type_combo.pack(padx=5, pady=(0, 15))
    type_combo.set(types[0] if types else "")

    # Поле "Уровень"
    tk.Label(main_frame, text="Уровень:", font=("Arial", 10, "bold"),
             bg="#f0f0f0", anchor="w").pack(fill="x", padx=5, pady=(0, 5))
    level_combo = ttk.Combobox(main_frame,
                               values=["Локальный", "Региональный", "Национальный", "Международный"],
                               state="readonly", font=("Arial", 10), width=58)
    level_combo.pack(padx=5, pady=(0, 15))
    level_combo.set("Локальный")

    # Поле "Описание"
    tk.Label(main_frame, text="Описание:", font=("Arial", 10, "bold"),
             bg="#f0f0f0", anchor="w").pack(fill="x", padx=5, pady=(0, 5))
    desc_frame = tk.Frame(main_frame)
    desc_frame.pack(padx=5, pady=(0, 15), fill="both")

    desc_text = tk.Text(desc_frame, height=6, width=60, font=("Arial", 10),
                        relief="solid", bd=1, wrap="word")
    scrollbar = tk.Scrollbar(desc_frame, command=desc_text.yview)
    desc_text.config(yscrollcommand=scrollbar.set)

    desc_text.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Кнопка "Сохранить"
    btn_frame = tk.Frame(main_frame)
    btn_frame.pack(pady=20)

    save_btn = tk.Button(btn_frame, text="Сохранить достижение",
                         font=("Arial", 11, "bold"),
                         bg="#4CAF50", fg="white", relief="raised",
                         padx=20, pady=8, cursor="hand2")
    save_btn.pack()

    return name_entry, date_entry, type_combo, level_combo, desc_text, save_btn


def create_list_tab(parent):
    """Создание вкладки со списком достижений"""
    main_frame = ttk.Frame(parent)
    main_frame.pack(padx=20, pady=20, fill="both", expand=True)

    # Заголовок
    header_frame = tk.Frame(main_frame, bg="#2c3e50")
    header_frame.pack(fill="x", pady=(0, 15))

    tk.Label(header_frame, text="Мои учебные достижения",
             font=("Arial", 14, "bold"), bg="#2c3e50", fg="white",
             padx=10, pady=10).pack()

    # Фрейм для списка с прокруткой
    list_frame = tk.Frame(main_frame)
    list_frame.pack(fill="both", expand=True, pady=(0, 15))

    # Создаем Treeview для отображения в табличном формате
    columns = ("Дата", "Название", "Тип", "Уровень")
    tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=12)

    # Настройка колонок
    tree.heading("Дата", text="Дата", anchor="w")
    tree.heading("Название", text="Название", anchor="w")
    tree.heading("Тип", text="Тип", anchor="w")
    tree.heading("Уровень", text="Уровень", anchor="w")

    tree.column("Дата", width=100)
    tree.column("Название", width=250)
    tree.column("Тип", width=120)
    tree.column("Уровень", width=120)

    # Добавляем прокрутку
    scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar.set)

    tree.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Фрейм для кнопок
    button_frame = tk.Frame(main_frame)
    button_frame.pack(fill="x", pady=10)

    # Кнопка обновления
    refresh_btn = tk.Button(button_frame, text="Обновить список",
                            font=("Arial", 10, "bold"),
                            bg="#3498db", fg="white", relief="raised",
                            padx=15, pady=6, cursor="hand2")
    refresh_btn.pack(side="left", padx=(0, 10))

    # Кнопка удаления
    delete_btn = tk.Button(button_frame, text="Удалить выбранное",
                           font=("Arial", 10, "bold"),
                           bg="#e74c3c", fg="white", relief="raised",
                           padx=15, pady=6, cursor="hand2")
    delete_btn.pack(side="left", padx=(0, 10))

    # Кнопка экспорта
    export_btn = tk.Button(button_frame, text="Экспорт в Word",
                           font=("Arial", 10, "bold"),
                           bg="#9b59b6", fg="white", relief="raised",
                           padx=15, pady=6, cursor="hand2")
    export_btn.pack(side="left")

    return tree, refresh_btn, delete_btn, export_btn


def refresh_treeview(tree):
    """Обновление Treeview данными из БД"""
    # Очищаем текущие записи
    for item in tree.get_children():
        tree.delete(item)

    # Загружаем новые данные
    records = load_records()
    for record in records:
        tree.insert("", "end", values=record)


def on_save(name_entry, date_entry, type_combo, level_combo, desc_text, tree):
    """Обработчик сохранения записи"""
    name = name_entry.get().strip()
    date = date_entry.get().strip()
    typ = type_combo.get()
    level = level_combo.get()
    desc = desc_text.get("1.0", "end-1c").strip()

    if not name:
        messagebox.showwarning("Внимание", "Пожалуйста, введите название достижения")
        name_entry.focus_set()
        return

    if not date:
        messagebox.showwarning("Внимание", "Пожалуйста, введите дату в формате ГГГГ-ММ-ДД")
        date_entry.focus_set()
        return

    # Простая проверка формата даты
    if len(date) != 10 or date[4] != '-' or date[7] != '-':
        messagebox.showwarning("Внимание", "Дата должна быть в формате ГГГГ-ММ-ДД")
        date_entry.focus_set()
        return

    try:
        save_to_db(name, date, typ, level, desc)

        # Очищаем поля формы
        name_entry.delete(0, tk.END)
        date_entry.delete(0, tk.END)
        desc_text.delete("1.0", tk.END)

        # Обновляем список достижений
        refresh_treeview(tree)

        messagebox.showinfo("Успех", f"Достижение '{name}' успешно сохранено!")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить данные: {e}")


def on_delete(tree):
    """Обработчик удаления записи"""
    selected_item = tree.selection()
    if not selected_item:
        messagebox.showwarning("Внимание", "Выберите запись для удаления")
        return

    # Получаем данные выбранной записи
    item_values = tree.item(selected_item[0], "values")

    # Подтверждение удаления
    if messagebox.askyesno("Подтверждение",
                           f"Вы уверены, что хотите удалить достижение:\n{item_values[1]}?"):

        # Получаем индекс записи
        all_items = tree.get_children()
        selected_index = all_items.index(selected_item[0])

        if delete_record(selected_index):
            # Удаляем из Treeview
            tree.delete(selected_item[0])
            messagebox.showinfo("Успех", "Запись успешно удалена")
        else:
            messagebox.showerror("Ошибка", "Не удалось удалить запись")


def export_to_word():
    """Экспорт данных в Word документ"""
    try:
        doc = Document()

        # Заголовок
        title = doc.add_heading("Личные учебные достижения", 0)
        title.alignment = 1  # Центрирование

        # Подзаголовок
        doc.add_paragraph(f"Отчет сформирован: {get_current_date()}")
        doc.add_paragraph()

        # Данные
        records = load_records_with_desc()

        if not records:
            doc.add_paragraph("Нет сохраненных достижений.")
        else:
            for i, (date, name, typ, level, desc) in enumerate(records, 1):
                # Добавляем номер и основную информацию
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(name).bold = True
                p.add_run(f" — {date}").italic = True
                p.add_run(f" ({typ}, {level})")

                # Добавляем описание, если есть
                if desc:
                    desc_para = doc.add_paragraph()
                    desc_para.add_run("Описание: ").italic = True
                    desc_para.add_run(desc)

                doc.add_paragraph()  # Пустая строка между записями

        # Сохраняем документ
        filename = f"достижения_{get_current_date()}.docx"
        doc.save(filename)

        messagebox.showinfo("Экспорт завершен",
                            f"Отчет успешно сохранен в файл:\n{filename}")

    except Exception as e:
        messagebox.showerror("Ошибка экспорта", f"Не удалось создать документ: {e}")


def get_current_date():
    """Получение текущей даты в формате ГГГГ-ММ-ДД"""
    from datetime import datetime
    return datetime.now().strftime("%Y-%m-%d")


# ========== ОСНОВНАЯ ЧАСТЬ ПРОГРАММЫ ==========

def main():
    """Основная функция программы"""
    # Создание главного окна
    root = tk.Tk()
    root.title("Журнал личных учебных достижений")
    root.geometry("800x600")
    root.resizable(True, True)

    # Устанавливаем иконку (если есть)
    try:
        root.iconbitmap(default="icon.ico")
    except:
        pass

    # Создаем стиль для Notebook
    style = ttk.Style()
    style.configure("TNotebook.Tab", font=("Arial", 11, "bold"), padding=[10, 5])

    # Создаем Notebook (вкладки)
    notebook = ttk.Notebook(root)
    notebook.pack(fill="both", expand=True, padx=10, pady=10)

    # Вкладка "Добавить достижение"
    tab_add = ttk.Frame(notebook)
    notebook.add(tab_add, text="➕ Добавить достижение")

    # Вкладка "Мои достижения"
    tab_list = ttk.Frame(notebook)
    notebook.add(tab_list, text="📋 Мои достижения")

    # Создаем форму добавления
    name_entry, date_entry, type_combo, level_combo, desc_text, save_btn = create_add_form(tab_add)

    # Создаем список достижений
    tree, refresh_btn, delete_btn, export_btn = create_list_tab(tab_list)

    # Привязываем обработчики событий
    save_btn.config(command=lambda: on_save(name_entry, date_entry, type_combo, level_combo, desc_text, tree))
    refresh_btn.config(command=lambda: refresh_treeview(tree))
    delete_btn.config(command=lambda: on_delete(tree))
    export_btn.config(command=export_to_word)

    # Инициализируем базу данных
    init_db()

    # Загружаем начальные данные
    refresh_treeview(tree)

    # Запускаем главный цикл
    root.mainloop()


if __name__ == "__main__":
    main()