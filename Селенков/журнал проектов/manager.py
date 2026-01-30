import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import psycopg2
from psycopg2 import Error
import os
from datetime import datetime, timedelta
import markdown
import openpyxl
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image as XLImage
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import sys
import traceback
from PIL import Image as PILImage

# Конфигурация базы данных
DB_CONFIG = {
    'host': 'localhost',
    'database': 'postgres',
    'user': 'postgres',
    'password': '1111',
    'port': '5432'
}


class ProjectManagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Менеджер проектов")
        self.root.geometry("1300x850")
        self.root.configure(bg='#f0f0f0')

        # Устанавливаем стиль
        self.setup_styles()

        # Текущий выбранный проект
        self.current_project_id = None
        self.current_project_file = None

        # Словарь для хранения технологий проекта
        self.project_technologies = {}

        # Создаем структуру БД при запуске
        self.init_database()

        # Создаем папки для хранения данных
        self.create_folders()

        # Строим интерфейс
        self.setup_ui()

        # Загружаем проекты
        self.load_projects()

        # Обработка закрытия окна
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def setup_styles(self):
        """Настройка стилей для интерфейса"""
        style = ttk.Style()
        style.theme_use('clam')

        # Кастомные стили
        style.configure('Title.TLabel', font=('Segoe UI', 14, 'bold'))
        style.configure('Header.TLabel', font=('Segoe UI', 11, 'bold'))
        style.configure('Custom.TButton', font=('Segoe UI', 10))
        style.configure('Success.TLabel', font=('Segoe UI', 10), foreground='green')

        # Стиль для Treeview
        style.configure("Treeview.Heading", font=('Segoe UI', 10, 'bold'))
        style.configure("Treeview", font=('Segoe UI', 10), rowheight=25)

    def create_folders(self):
        """Создание необходимых папок"""
        folders = ['projects', 'reports', 'reports/charts']
        for folder in folders:
            if not os.path.exists(folder):
                os.makedirs(folder, exist_ok=True)

    def init_database(self):
        """Инициализация структуры базы данных"""
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Таблица проектов
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS projects (
                    id SERIAL PRIMARY KEY,
                    name VARCHAR(255) NOT NULL,
                    discipline VARCHAR(255),
                    status VARCHAR(100),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    file_path TEXT
                )
            """)

            # Таблица технологий
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS technologies (
                    id SERIAL PRIMARY KEY,
                    project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
                    technology VARCHAR(255) NOT NULL,
                    added_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Таблица логов действий
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS activity_log (
                    id SERIAL PRIMARY KEY,
                    project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
                    action_type VARCHAR(50) NOT NULL,
                    action_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    details TEXT
                )
            """)

            conn.commit()
            cursor.close()
            conn.close()

        except Error as e:
            messagebox.showerror("Ошибка БД", f"Не удалось инициализировать БД:\n{str(e)}")
            self.root.destroy()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Основной контейнер
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)

        # Заголовок
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=tk.X, pady=(0, 15))

        ttk.Label(title_frame, text="📋 МЕНЕДЖЕР ПРОЕКТОВ",
                  style='Title.TLabel').pack(side=tk.LEFT)

        ttk.Label(title_frame, text=f"Версия 1.0 | {datetime.now().strftime('%d.%m.%Y')}",
                  foreground='gray').pack(side=tk.RIGHT)

        # 1. Панель управления
        control_frame = ttk.LabelFrame(main_container, text="Панель управления",
                                       padding=15)
        control_frame.pack(fill=tk.X, pady=(0, 15))

        # Ввод данных проекта
        input_frame = ttk.Frame(control_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))

        # Название проекта
        ttk.Label(input_frame, text="Название проекта:", width=15,
                  anchor='e').grid(row=0, column=0, padx=5, pady=5, sticky='e')
        self.project_name_entry = ttk.Entry(input_frame, width=35, font=('Segoe UI', 10))
        self.project_name_entry.grid(row=0, column=1, padx=5, pady=5, sticky='w')

        # Дисциплина
        ttk.Label(input_frame, text="Дисциплина:", width=15,
                  anchor='e').grid(row=0, column=2, padx=20, pady=5, sticky='e')
        self.discipline_entry = ttk.Entry(input_frame, width=25, font=('Segoe UI', 10))
        self.discipline_entry.grid(row=0, column=3, padx=5, pady=5, sticky='w')

        # Статус
        ttk.Label(input_frame, text="Статус:", width=15,
                  anchor='e').grid(row=0, column=4, padx=20, pady=5, sticky='e')
        self.status_combobox = ttk.Combobox(input_frame, width=20, font=('Segoe UI', 10),
                                            values=["В процессе", "Завершен", "На паузе", "Планируется"])
        self.status_combobox.grid(row=0, column=5, padx=5, pady=5, sticky='w')
        self.status_combobox.set("В процессе")

        # Кнопки управления
        buttons_frame = ttk.Frame(control_frame)
        buttons_frame.pack(fill=tk.X)

        button_configs = [
            ("➕ Создать", self.create_project, '#4CAF50'),
            ("💾 Сохранить", self.save_project, '#2196F3'),
            ("🗑️ Удалить", self.delete_project, '#F44336'),
            ("📄 Открыть", self.open_description, '#FF9800'),
            ("📊 Excel", self.export_to_excel, '#009688'),
            ("📝 Word", self.export_to_word, '#673AB7')
        ]

        for text, command, color in button_configs:
            btn = tk.Button(buttons_frame, text=text, command=command,
                            bg=color, fg='white', font=('Segoe UI', 10, 'bold'),
                            padx=15, pady=8, bd=0, cursor='hand2')
            btn.pack(side=tk.LEFT, padx=5)
            btn.bind("<Enter>", lambda e, b=btn: b.configure(bg='#333333'))
            btn.bind("<Leave>", lambda e, b=btn, c=color: b.configure(bg=c))

        # 2. Основная рабочая область
        work_area = ttk.Frame(main_container)
        work_area.pack(fill=tk.BOTH, expand=True)

        # Список проектов слева (40%)
        list_frame = ttk.LabelFrame(work_area, text="📂 Список проектов", padding=10)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))

        # Панель поиска/фильтрации
        search_frame = ttk.Frame(list_frame)
        search_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(search_frame, text="Поиск:").pack(side=tk.LEFT, padx=(0, 5))
        self.search_entry = ttk.Entry(search_frame, width=30)
        self.search_entry.pack(side=tk.LEFT, padx=(0, 10))
        self.search_entry.bind('<KeyRelease>', self.filter_projects)

        ttk.Button(search_frame, text="Обновить",
                   command=self.load_projects, width=10).pack(side=tk.RIGHT)

        # Treeview для проектов
        columns = ("Название", "Дисциплина", "Статус", "Создан", "Обновлен")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)

        # Настройка колонок
        col_widths = [200, 120, 100, 120, 120]
        for idx, col in enumerate(columns):
            self.tree.heading(col, text=col,
                              command=lambda c=col: self.sort_treeview(c))
            self.tree.column(col, width=col_widths[idx])

        # Настройка тэгов для разных статусов
        self.tree.tag_configure('active', background='#e8f5e9')
        self.tree.tag_configure('completed', background='#e3f2fd')
        self.tree.tag_configure('paused', background='#fff3e0')
        self.tree.tag_configure('planned', background='#f5f5f5')

        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Привязка события выбора
        self.tree.bind("<<TreeviewSelect>>", self.on_project_select)
        self.tree.bind("<Double-1>", lambda e: self.open_description())

        # 3. Область редактирования справа (60%)
        editor_frame = ttk.LabelFrame(work_area, text="✏️ Редактирование проекта", padding=10)
        editor_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Панель инструментов редактора
        toolbar = ttk.Frame(editor_frame)
        toolbar.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(toolbar, text="Форматирование: ").pack(side=tk.LEFT)

        # Кнопки форматирования
        formats = [
            ("Ж", lambda: self.insert_text("**текст**")),
            ("К", lambda: self.insert_text("*текст*")),
            ("H1", lambda: self.insert_text("# Заголовок 1")),
            ("H2", lambda: self.insert_text("## Заголовок 2")),
            ("•", lambda: self.insert_text("- пункт списка")),
            ("1.", lambda: self.insert_text("1. нумерованный пункт"))
        ]

        for text, command in formats:
            btn = ttk.Button(toolbar, text=text, width=3, command=command)
            btn.pack(side=tk.LEFT, padx=2)

        ttk.Button(toolbar, text="Предпросмотр",
                   command=self.preview_markdown).pack(side=tk.RIGHT)

        # Текстовый редактор
        text_container = ttk.Frame(editor_frame)
        text_container.pack(fill=tk.BOTH, expand=True)

        self.text_editor = tk.Text(text_container, wrap=tk.WORD,
                                   font=("Consolas", 11),
                                   bg='white', relief=tk.SUNKEN,
                                   padx=10, pady=10)

        text_scrollbar = ttk.Scrollbar(text_container, orient=tk.VERTICAL,
                                       command=self.text_editor.yview)
        self.text_editor.configure(yscrollcommand=text_scrollbar.set)

        self.text_editor.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 4. Панель технологий
        tech_frame = ttk.LabelFrame(main_container, text="🛠️ Технологии проекта",
                                    padding=10)
        tech_frame.pack(fill=tk.X, pady=(15, 0))

        # Ввод новой технологии
        input_tech_frame = ttk.Frame(tech_frame)
        input_tech_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(input_tech_frame, text="Добавить технологию:").pack(side=tk.LEFT, padx=(0, 5))
        self.tech_entry = ttk.Entry(input_tech_frame, width=30)
        self.tech_entry.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(input_tech_frame, text="Добавить",
                   command=self.add_technology,
                   style='Custom.TButton').pack(side=tk.LEFT)

        self.tech_entry.bind('<Return>', lambda e: self.add_technology())

        # Область отображения технологий
        self.tech_display_canvas = tk.Canvas(tech_frame, height=100, bg='white',
                                             highlightthickness=1,
                                             highlightbackground='#ddd')
        self.tech_display_canvas.pack(fill=tk.X)

        self.tech_scrollbar = ttk.Scrollbar(tech_frame, orient=tk.HORIZONTAL,
                                            command=self.tech_display_canvas.xview)
        self.tech_display_canvas.configure(xscrollcommand=self.tech_scrollbar.set)
        self.tech_scrollbar.pack(fill=tk.X)

        self.tech_inner_frame = ttk.Frame(self.tech_display_canvas)
        self.tech_window = self.tech_display_canvas.create_window(
            (0, 0), window=self.tech_inner_frame, anchor='nw')

        # 5. Вкладка аналитики
        self.setup_analytics_tab(main_container)

        # Статус бар
        self.status_bar = ttk.Label(main_container, text="Готово",
                                    relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(fill=tk.X, pady=(10, 0))

    def setup_analytics_tab(self, parent):
        """Настройка вкладки аналитики"""
        analytics_frame = ttk.LabelFrame(parent, text="📈 Аналитика и отчётность",
                                         padding=15)
        analytics_frame.pack(fill=tk.X, pady=(15, 0))

        # Кнопки отчетов
        report_buttons_frame = ttk.Frame(analytics_frame)
        report_buttons_frame.pack(fill=tk.X)

        ttk.Button(report_buttons_frame, text="📊 Сформировать комплексный отчёт",
                   command=self.generate_report,
                   style='Custom.TButton').pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(report_buttons_frame, text="📈 Быстрая статистика",
                   command=self.show_quick_stats).pack(side=tk.LEFT)

        # Область информации
        self.report_info_frame = ttk.Frame(analytics_frame)
        self.report_info_frame.pack(fill=tk.X, pady=(10, 0))

        self.report_info_label = ttk.Label(self.report_info_frame,
                                           text="",
                                           style='Success.TLabel')
        self.report_info_label.pack()

    def insert_text(self, text):
        """Вставка форматированного текста"""
        if self.text_editor.tag_ranges(tk.SEL):
            # Заменяем выделенный текст
            self.text_editor.delete(tk.SEL_FIRST, tk.SEL_LAST)
            self.text_editor.insert(tk.SEL_FIRST, text)
        else:
            # Вставляем в позицию курсора
            self.text_editor.insert(tk.INSERT, text)

    def preview_markdown(self):
        """Предпросмотр Markdown"""
        if not self.text_editor.get(1.0, tk.END).strip():
            messagebox.showinfo("Предпросмотр", "Текст для предпросмотра отсутствует")
            return

        preview_window = tk.Toplevel(self.root)
        preview_window.title("Предпросмотр Markdown")
        preview_window.geometry("800x600")

        text = self.text_editor.get(1.0, tk.END)
        html = markdown.markdown(text, extensions=['extra'])

        # Простой предпросмотр
        text_widget = tk.Text(preview_window, wrap=tk.WORD, font=("Arial", 11))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Вставляем текст с базовым форматированием
        lines = text.split('\n')
        for line in lines:
            if line.startswith('# '):
                text_widget.insert(tk.END, line[2:] + '\n', 'h1')
            elif line.startswith('## '):
                text_widget.insert(tk.END, line[3:] + '\n', 'h2')
            elif line.startswith('### '):
                text_widget.insert(tk.END, line[4:] + '\n', 'h3')
            elif line.startswith('**') and line.endswith('**'):
                text_widget.insert(tk.END, line[2:-2] + '\n', 'bold')
            elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                text_widget.insert(tk.END, line[1:-1] + '\n', 'italic')
            else:
                text_widget.insert(tk.END, line + '\n')

        # Настраиваем тэги
        text_widget.tag_config('h1', font=('Arial', 16, 'bold'))
        text_widget.tag_config('h2', font=('Arial', 14, 'bold'))
        text_widget.tag_config('h3', font=('Arial', 12, 'bold'))
        text_widget.tag_config('bold', font=('Arial', 11, 'bold'))
        text_widget.tag_config('italic', font=('Arial', 11, 'italic'))

        text_widget.config(state=tk.DISABLED)

    def filter_projects(self, event=None):
        """Фильтрация проектов по поисковому запросу"""
        search_term = self.search_entry.get().lower()

        for item in self.tree.get_children():
            values = self.tree.item(item)['values']
            if search_term in ' '.join(str(v).lower() for v in values):
                self.tree.item(item, tags=())
            else:
                self.tree.item(item, tags=('hidden',))

        self.tree.tag_configure('hidden', foreground='gray')

    def log_activity(self, project_id, action_type, details=""):
        """Логирование действий с параметризованными запросами"""
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                INSERT INTO activity_log (project_id, action_type, details)
                VALUES (%s, %s, %s)
            """, (project_id, action_type, details))

            conn.commit()
            cursor.close()
            conn.close()

        except Error as e:
            print(f"Ошибка логирования: {e}")
            self.status_bar.config(text=f"Ошибка логирования: {e}")

    def create_project(self):
        """Создание нового проекта"""
        name = self.project_name_entry.get().strip()
        discipline = self.discipline_entry.get().strip()
        status = self.status_combobox.get()

        if not name:
            messagebox.showwarning("Предупреждение", "Введите название проекта")
            self.project_name_entry.focus()
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Проверяем, существует ли проект с таким именем
            cursor.execute("SELECT id FROM projects WHERE name = %s", (name,))
            if cursor.fetchone():
                messagebox.showwarning("Предупреждение",
                                       f"Проект с именем '{name}' уже существует")
                return

            # Создаем файл для проекта
            safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            file_name = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            file_path = os.path.join('projects', file_name)

            # Создаем файл с шаблоном
            template = f"""# {name}

## Описание проекта
*Здесь будет описание вашего проекта*

## Цели проекта
- Цель 1
- Цель 2

## Задачи
- [ ] Задача 1
- [ ] Задача 2

## Результаты
*Ожидаемые результаты*
"""

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(template)

            # Сохраняем в БД с параметризованным запросом
            cursor.execute("""
                INSERT INTO projects (name, discipline, status, file_path)
                VALUES (%s, %s, %s, %s) RETURNING id
            """, (name, discipline, status, file_path))

            project_id = cursor.fetchone()[0]

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(project_id, "CREATE", f"Создан проект: {name}")

            # Обновляем список
            self.load_projects()

            # Очищаем поля
            self.project_name_entry.delete(0, tk.END)
            self.discipline_entry.delete(0, tk.END)

            self.status_bar.config(text=f"Проект '{name}' успешно создан")
            messagebox.showinfo("Успех", f"Проект '{name}' успешно создан!")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось создать проект:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Неожиданная ошибка:\n{str(e)}")
            self.status_bar.config(text=f"Неожиданная ошибка: {str(e)}")

    def load_projects(self):
        """Загрузка проектов из БД в Treeview"""
        # Очищаем текущий список
        for item in self.tree.get_children():
            self.tree.delete(item)

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT id, name, discipline, status, 
                       created_at, updated_at 
                FROM projects ORDER BY updated_at DESC
            """)

            projects = cursor.fetchall()

            for project in projects:
                # Форматируем даты
                created = project[4].strftime('%d.%m.%Y') if project[4] else ''
                updated = project[5].strftime('%d.%m.%Y') if project[5] else ''

                # Определяем тэг в зависимости от статуса
                status = project[3]
                tag = ''
                if status == 'В процессе':
                    tag = 'active'
                elif status == 'Завершен':
                    tag = 'completed'
                elif status == 'На паузе':
                    tag = 'paused'
                elif status == 'Планируется':
                    tag = 'planned'

                self.tree.insert("", tk.END, values=(
                    project[1],  # name
                    project[2],  # discipline
                    project[3],  # status
                    created,
                    updated
                ), tags=(tag, f"id_{project[0]}"))

            cursor.close()
            conn.close()

            self.status_bar.config(text=f"Загружено проектов: {len(projects)}")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить проекты:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка загрузки: {str(e)}")

    def on_project_select(self, event):
        """Обработка выбора проекта"""
        try:
            selection = self.tree.selection()
            if not selection:
                return

            item = self.tree.item(selection[0])
            project_name = item['values'][0]

            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Получаем ID проекта и путь к файлу с параметризованным запросом
            cursor.execute("""
                SELECT id, file_path FROM projects WHERE name = %s
            """, (project_name,))

            result = cursor.fetchone()
            if result:
                self.current_project_id = result[0]
                self.current_project_file = result[1]

                # Загружаем описание из файла
                if os.path.exists(self.current_project_file):
                    try:
                        with open(self.current_project_file, 'r', encoding='utf-8') as f:
                            content = f.read()

                        self.text_editor.delete(1.0, tk.END)
                        self.text_editor.insert(1.0, content)
                    except UnicodeDecodeError:
                        # Пробуем другую кодировку
                        with open(self.current_project_file, 'r', encoding='cp1251') as f:
                            content = f.read()
                        self.text_editor.delete(1.0, tk.END)
                        self.text_editor.insert(1.0, content)

                # Загружаем технологии
                self.load_technologies()

                # Обновляем поля ввода
                cursor.execute("""
                    SELECT name, discipline, status FROM projects WHERE id = %s
                """, (self.current_project_id,))

                proj_data = cursor.fetchone()
                if proj_data:
                    self.project_name_entry.delete(0, tk.END)
                    self.project_name_entry.insert(0, proj_data[0])
                    self.discipline_entry.delete(0, tk.END)
                    self.discipline_entry.insert(0, proj_data[1])
                    self.status_combobox.set(proj_data[2])

                self.status_bar.config(text=f"Выбран проект: {project_name}")

            cursor.close()
            conn.close()

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить проект:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка чтения файла:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка файла: {str(e)}")

    def save_project(self):
        """Сохранение изменений проекта"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для сохранения")
            return

        name = self.project_name_entry.get().strip()
        discipline = self.discipline_entry.get().strip()
        status = self.status_combobox.get()

        if not name:
            messagebox.showwarning("Предупреждение", "Введите название проекта")
            self.project_name_entry.focus()
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Обновляем данные в БД с параметризованным запросом
            cursor.execute("""
                UPDATE projects 
                SET name = %s, discipline = %s, status = %s, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
            """, (name, discipline, status, self.current_project_id))

            # Сохраняем описание в файл
            content = self.text_editor.get(1.0, tk.END)
            if self.current_project_file:
                # Создаем резервную копию
                if os.path.exists(self.current_project_file):
                    backup_path = self.current_project_file + '.backup'
                    os.replace(self.current_project_file, backup_path)

                with open(self.current_project_file, 'w', encoding='utf-8') as f:
                    f.write(content)

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "UPDATE", f"Обновлен проект: {name}")

            # Обновляем список
            self.load_projects()

            self.status_bar.config(text=f"Проект '{name}' успешно сохранен")
            messagebox.showinfo("Успех", "Проект успешно сохранен!")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить проект:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка сохранения: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка записи файла:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка файла: {str(e)}")

    def delete_project(self):
        """Удаление проекта"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для удаления")
            return

        project_name = self.project_name_entry.get().strip()

        if not messagebox.askyesno("Подтверждение",
                                   f"Вы уверены, что хотите удалить проект '{project_name}'?"):
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Получаем путь к файлу
            cursor.execute("SELECT file_path FROM projects WHERE id = %s",
                           (self.current_project_id,))
            file_path = cursor.fetchone()[0]

            # Удаляем из БД (каскадно удалятся технологии и логи)
            cursor.execute("DELETE FROM projects WHERE id = %s",
                           (self.current_project_id,))

            # Удаляем файл с проверкой существования
            if file_path and os.path.exists(file_path):
                # Перемещаем в корзину вместо полного удаления
                trash_dir = 'trash'
                if not os.path.exists(trash_dir):
                    os.makedirs(trash_dir, exist_ok=True)

                trash_path = os.path.join(trash_dir,
                                          os.path.basename(file_path) +
                                          f".deleted_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
                os.rename(file_path, trash_path)

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(None, "DELETE", f"Удален проект: {project_name}")

            # Очищаем интерфейс
            self.current_project_id = None
            self.current_project_file = None
            self.text_editor.delete(1.0, tk.END)
            self.project_name_entry.delete(0, tk.END)
            self.discipline_entry.delete(0, tk.END)
            self.clear_technologies_display()

            # Обновляем список
            self.load_projects()

            self.status_bar.config(text=f"Проект '{project_name}' удален")
            messagebox.showinfo("Успех", f"Проект '{project_name}' удален")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить проект:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка удаления: {str(e)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при удалении:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка: {str(e)}")

    def open_description(self):
        """Открытие описания во внешнем просмотрщике"""
        if not self.current_project_file or not os.path.exists(self.current_project_file):
            messagebox.showwarning("Предупреждение", "Файл описания не найден")
            return

        try:
            # Пытаемся открыть файл в системном просмотрщике
            if sys.platform == "win32":
                os.startfile(self.current_project_file)
            elif sys.platform == "darwin":
                os.system(f"open '{self.current_project_file}'")
            else:
                os.system(f"xdg-open '{self.current_project_file}'")

            self.status_bar.config(text="Файл открыт во внешнем приложении")
        except:
            messagebox.showinfo("Информация",
                                f"Файл расположен по пути:\n{os.path.abspath(self.current_project_file)}")

    def add_technology(self):
        """Добавление технологии к проекту"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение",
                                   "Выберите проект для добавления технологии")
            return

        tech = self.tech_entry.get().strip()
        if not tech:
            messagebox.showwarning("Предупреждение", "Введите название технологии")
            self.tech_entry.focus()
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # Проверяем, есть ли уже такая технология у проекта
            cursor.execute("""
                SELECT id FROM technologies 
                WHERE project_id = %s AND LOWER(technology) = LOWER(%s)
            """, (self.current_project_id, tech))

            if cursor.fetchone():
                messagebox.showwarning("Предупреждение",
                                       "Эта технология уже добавлена к проекту")
                return

            # Добавляем технологию с параметризованным запросом
            cursor.execute("""
                INSERT INTO technologies (project_id, technology)
                VALUES (%s, %s)
            """, (self.current_project_id, tech))

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "ADD_TECH",
                              f"Добавлена технология: {tech}")

            # Обновляем отображение
            self.load_technologies()

            # Очищаем поле ввода
            self.tech_entry.delete(0, tk.END)

            self.status_bar.config(text=f"Технология '{tech}' добавлена")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось добавить технологию:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка: {str(e)}")

    def load_technologies(self):
        """Загрузка технологий проекта"""
        if not self.current_project_id:
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT technology FROM technologies 
                WHERE project_id = %s 
                ORDER BY added_at DESC
            """, (self.current_project_id,))

            technologies = [row[0] for row in cursor.fetchall()]

            cursor.close()
            conn.close()

            # Отображаем технологии
            self.display_technologies(technologies)

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить технологии:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка загрузки технологий: {str(e)}")

    def display_technologies(self, technologies):
        """Отображение технологий в интерфейсе"""
        # Очищаем текущее отображение
        self.clear_technologies_display()

        if not technologies:
            label = ttk.Label(self.tech_inner_frame, text="Технологии не добавлены",
                              foreground="gray", font=('Segoe UI', 10))
            label.pack(pady=20)
            return

        # Создаем фреймы для технологий
        for tech in technologies:
            tech_frame = ttk.Frame(self.tech_inner_frame, relief=tk.RAISED)
            tech_frame.pack(side=tk.LEFT, padx=5, pady=5, fill=tk.Y)

            label = tk.Label(tech_frame, text=tech, bg='#e3f2fd',
                             fg='#1565c0', font=('Segoe UI', 9, 'bold'),
                             padx=10, pady=5, relief=tk.RIDGE)
            label.pack(side=tk.LEFT, padx=(0, 5))

            # Кнопка удаления
            btn = tk.Button(tech_frame, text="×", command=lambda t=tech: self.remove_technology(t),
                            bg='#ff5252', fg='white', font=('Arial', 10, 'bold'),
                            width=2, height=1, bd=0, cursor='hand2')
            btn.pack(side=tk.RIGHT, padx=(0, 2))
            btn.bind("<Enter>", lambda e, b=btn: b.configure(bg='#ff0000'))
            btn.bind("<Leave>", lambda e, b=btn: b.configure(bg='#ff5252'))

        # Обновляем область прокрутки
        self.tech_inner_frame.update_idletasks()
        self.tech_display_canvas.config(scrollregion=self.tech_display_canvas.bbox("all"))

    def clear_technologies_display(self):
        """Очистка области отображения технологий"""
        for widget in self.tech_inner_frame.winfo_children():
            widget.destroy()

    def remove_technology(self, technology):
        """Удаление технологии из проекта"""
        if not self.current_project_id:
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                DELETE FROM technologies 
                WHERE project_id = %s AND technology = %s
            """, (self.current_project_id, technology))

            conn.commit()
            cursor.close()
            conn.close()

            # Логируем действие
            self.log_activity(self.current_project_id, "REMOVE_TECH",
                              f"Удалена технология: {technology}")

            # Обновляем отображение
            self.load_technologies()

            self.status_bar.config(text=f"Технология '{technology}' удалена")

        except Error as e:
            messagebox.showerror("Ошибка", f"Не удалось удалить технологию:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка удаления: {str(e)}")

    def sort_treeview(self, column):
        """Сортировка Treeview по колонке"""
        # Получаем текущие данные
        items = [(self.tree.set(child, column), child) for child in self.tree.get_children('')]

        # Определяем тип сортировки
        try:
            # Пытаемся отсортировать как дату
            items.sort(key=lambda x: datetime.strptime(x[0], '%d.%m.%Y')
            if x[0] and '.' in x[0] else x[0])
        except:
            # Сортируем как строку
            items.sort(key=lambda x: x[0].lower() if x[0] else '')

        # Перестраиваем Treeview
        for index, (_, child) in enumerate(items):
            self.tree.move(child, '', index)

    def generate_report(self):
        """Генерация комплексного отчета"""
        try:
            # Проверяем, есть ли данные
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM projects")
            count = cursor.fetchone()[0]
            cursor.close()
            conn.close()

            if count == 0:
                messagebox.showwarning("Нет данных", "Нет проектов для генерации отчета")
                return

            self.status_bar.config(text="Сбор данных для отчета...")
            self.root.update()

            # Собираем данные
            stats = self.collect_statistics()

            self.status_bar.config(text="Генерация Excel отчета...")
            self.root.update()

            # Генерируем Excel отчет
            excel_path = self.generate_excel_report(stats)

            self.status_bar.config(text="Генерация Word отчета...")
            self.root.update()

            # Генерируем Word отчет
            word_path = self.generate_word_report(stats)

            self.status_bar.config(text="Создание графиков...")
            self.root.update()

            # Создаем графики
            self.create_charts_for_reports(stats)

            # Обновляем информацию в интерфейсе
            self.report_info_label.config(
                text=f"✓ Отчеты сгенерированы:\n• {os.path.basename(excel_path)}\n• {os.path.basename(word_path)}"
            )

            # Показываем кнопки открытия
            open_frame = ttk.Frame(self.report_info_frame)
            open_frame.pack(pady=5)

            ttk.Button(open_frame, text="Открыть Excel",
                       command=lambda: self.open_file(excel_path)).pack(side=tk.LEFT, padx=5)
            ttk.Button(open_frame, text="Открыть Word",
                       command=lambda: self.open_file(word_path)).pack(side=tk.LEFT, padx=5)

            self.status_bar.config(text="Отчеты успешно сгенерированы!")
            messagebox.showinfo("Успех", "Отчеты успешно сгенерированы!")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сгенерировать отчет:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка генерации отчета: {str(e)}")
            print(traceback.format_exc())

    def open_file(self, filepath):
        """Открытие файла в системе"""
        try:
            if sys.platform == "win32":
                os.startfile(filepath)
            elif sys.platform == "darwin":
                os.system(f"open '{filepath}'")
            else:
                os.system(f"xdg-open '{filepath}'")
        except:
            messagebox.showinfo("Информация", f"Файл: {filepath}")

    def collect_statistics(self):
        """Сбор статистических данных с параметризованными запросами"""
        stats = {}

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            # 1. Количество проектов по дисциплинам
            cursor.execute("""
                SELECT COALESCE(discipline, 'Не указана'), COUNT(*) 
                FROM projects 
                GROUP BY discipline 
                ORDER BY COUNT(*) DESC
            """)
            stats['projects_by_discipline'] = dict(cursor.fetchall())

            # 2. Количество проектов по статусам
            cursor.execute("""
                SELECT COALESCE(status, 'Не указан'), COUNT(*) 
                FROM projects 
                GROUP BY status 
                ORDER BY COUNT(*) DESC
            """)
            stats['projects_by_status'] = dict(cursor.fetchall())

            # 3. Действия за последние 7 и 30 дней
            cursor.execute("""
                SELECT 
                    COUNT(CASE WHEN action_date >= CURRENT_DATE - INTERVAL '7 days' THEN 1 END) as last_7_days,
                    COUNT(CASE WHEN action_date >= CURRENT_DATE - INTERVAL '30 days' THEN 1 END) as last_30_days
                FROM activity_log
            """)
            actions = cursor.fetchone()
            stats['actions_last_7_days'] = actions[0] if actions else 0
            stats['actions_last_30_days'] = actions[1] if actions else 0

            # 4. Топ-5 самых часто используемых технологий
            cursor.execute("""
                SELECT technology, COUNT(*) as usage_count
                FROM technologies
                GROUP BY technology
                ORDER BY usage_count DESC
                LIMIT 5
            """)
            stats['top_technologies'] = dict(cursor.fetchall())

            # 5. Последние 5 проектов
            cursor.execute("""
                SELECT name, discipline, status, updated_at
                FROM projects
                ORDER BY updated_at DESC
                LIMIT 5
            """)
            stats['recent_projects'] = cursor.fetchall()

            # 6. Общее количество проектов
            cursor.execute("SELECT COUNT(*) FROM projects")
            stats['total_projects'] = cursor.fetchone()[0]

            # 7. Количество дисциплин
            stats['disciplines_count'] = len(stats['projects_by_discipline'])

            # 8. Общее количество технологий
            cursor.execute("SELECT COUNT(DISTINCT technology) FROM technologies")
            stats['unique_technologies'] = cursor.fetchone()[0]

            cursor.close()
            conn.close()

        except Error as e:
            raise Exception(f"Ошибка сбора статистики: {str(e)}")

        return stats

    def generate_excel_report(self, stats):
        """Генерация Excel отчета с графиками"""
        try:
            wb = Workbook()

            # Лист "Статистика"
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats['A1'] = "Отчет по проектам"
            ws_stats['A1'].font = openpyxl.styles.Font(bold=True, size=16, color="1F4E78")
            ws_stats.merge_cells('A1:E1')
            ws_stats['A1'].alignment = openpyxl.styles.Alignment(horizontal='center')

            ws_stats['A2'] = f"Сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
            ws_stats['A2'].font = openpyxl.styles.Font(italic=True, color="666666")
            ws_stats.merge_cells('A2:E2')
            ws_stats['A2'].alignment = openpyxl.styles.Alignment(horizontal='center')

            # Основные метрики
            ws_stats['A4'] = "Ключевые показатели"
            ws_stats['A4'].font = openpyxl.styles.Font(bold=True, size=12)

            data_rows = [
                ("Общее количество проектов", stats['total_projects']),
                ("Действий за 7 дней", stats['actions_last_7_days']),
                ("Действий за 30 дней", stats['actions_last_30_days']),
                ("Количество дисциплин", stats['disciplines_count']),
                ("Уникальных технологий", stats.get('unique_technologies', 0)),
            ]

            for i, (label, value) in enumerate(data_rows, start=5):
                ws_stats[f'A{i}'] = label
                ws_stats[f'A{i}'].font = openpyxl.styles.Font(bold=True)
                ws_stats[f'B{i}'] = value

                # Форматируем ячейки с данными
                cell = ws_stats[f'B{i}']
                cell.alignment = openpyxl.styles.Alignment(horizontal='right')
                if isinstance(value, (int, float)):
                    cell.number_format = '#,##0'

            # Проекты по дисциплинам
            start_row = len(data_rows) + 7
            ws_stats[f'A{start_row}'] = "Проекты по дисциплинам"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True, size=12)

            for i, (discipline, count) in enumerate(stats['projects_by_discipline'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = discipline
                ws_stats[f'B{start_row + i}'] = count

            # Проекты по статусам
            start_row += len(stats['projects_by_discipline']) + 3
            ws_stats[f'A{start_row}'] = "Проекты по статусам"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True, size=12)

            for i, (status, count) in enumerate(stats['projects_by_status'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = status
                ws_stats[f'B{start_row + i}'] = count

            # Топ технологий
            start_row += len(stats['projects_by_status']) + 3
            ws_stats[f'A{start_row}'] = "Топ-5 технологий"
            ws_stats[f'A{start_row}'].font = openpyxl.styles.Font(bold=True, size=12)

            for i, (tech, count) in enumerate(stats['top_technologies'].items(), start=1):
                ws_stats[f'A{start_row + i}'] = tech
                ws_stats[f'B{start_row + i}'] = count

            # Настраиваем ширину колонок
            for col in ['A', 'B']:
                ws_stats.column_dimensions[col].width = 25

            # Лист "Графики"
            ws_charts = wb.create_sheet("Графики")
            self.create_excel_charts(stats, ws_charts)

            # Сохраняем файл
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir, exist_ok=True)

            excel_path = os.path.join(report_dir, 'projects_report.xlsx')
            wb.save(excel_path)

            return os.path.abspath(excel_path)

        except Exception as e:
            print(f"Ошибка при генерации Excel отчета: {e}")
            # Пробуем альтернативный метод без графиков
            return self.generate_excel_simple(stats)

    def create_excel_charts(self, stats, ws_charts):
        """Создание встроенных графиков в Excel"""
        try:
            # 1. График по статусам
            ws_charts['A1'] = 'Статусы проектов'
            ws_charts['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws_charts.merge_cells('A1:B1')

            ws_charts['A3'] = 'Статус'
            ws_charts['B3'] = 'Количество'
            ws_charts['A3'].font = ws_charts['B3'].font = openpyxl.styles.Font(bold=True)

            row = 4
            for status, count in stats['projects_by_status'].items():
                ws_charts[f'A{row}'] = status
                ws_charts[f'B{row}'] = count
                row += 1

            chart1 = BarChart()
            chart1.type = "col"
            chart1.style = 10
            chart1.title = "Распределение проектов по статусам"
            chart1.y_axis.title = "Количество проектов"
            chart1.x_axis.title = "Статус"
            chart1.legend = None

            data1 = Reference(ws_charts, min_col=2, min_row=3, max_row=row - 1)
            cats1 = Reference(ws_charts, min_col=1, min_row=4, max_row=row - 1)
            chart1.add_data(data1, titles_from_data=True)
            chart1.set_categories(cats1)

            ws_charts.add_chart(chart1, "D2")

            # 2. График по дисциплинам
            ws_charts[f'A{row + 3}'] = 'Проекты по дисциплинам'
            ws_charts[f'A{row + 3}'].font = openpyxl.styles.Font(bold=True, size=14)
            ws_charts.merge_cells(f'A{row + 3}:B{row + 3}')

            ws_charts[f'A{row + 5}'] = 'Дисциплина'
            ws_charts[f'B{row + 5}'] = 'Количество'
            ws_charts[f'A{row + 5}'].font = ws_charts[f'B{row + 5}'].font = openpyxl.styles.Font(bold=True)

            row2 = row + 6
            for discipline, count in stats['projects_by_discipline'].items():
                ws_charts[f'A{row2}'] = discipline
                ws_charts[f'B{row2}'] = count
                row2 += 1

            chart2 = BarChart()
            chart2.type = "col"
            chart2.style = 11
            chart2.title = "Распределение проектов по дисциплинам"
            chart2.y_axis.title = "Количество проектов"
            chart2.x_axis.title = "Дисциплина"
            chart2.legend = None

            data2 = Reference(ws_charts, min_col=2, min_row=row + 5, max_row=row2 - 1)
            cats2 = Reference(ws_charts, min_col=1, min_row=row + 6, max_row=row2 - 1)
            chart2.add_data(data2, titles_from_data=True)
            chart2.set_categories(cats2)

            ws_charts.add_chart(chart2, "D20")

        except Exception as e:
            print(f"Ошибка при создании встроенных графиков: {e}")
            # Создаем графики как изображения
            self.create_excel_charts_as_images(stats, ws_charts)

    def create_excel_charts_as_images(self, stats, worksheet):
        """Создание графиков как изображений для Excel"""
        try:
            charts_dir = 'reports/charts'
            if not os.path.exists(charts_dir):
                os.makedirs(charts_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # 1. График проектов по статусам
            if stats['projects_by_status']:
                fig, ax = plt.subplots(figsize=(8, 5))
                labels = list(stats['projects_by_status'].keys())
                values = list(stats['projects_by_status'].values())

                colors = ['#4CAF50', '#2196F3', '#FF9800', '#F44336', '#9C27B0']
                bars = ax.bar(labels, values, color=colors[:len(labels)], edgecolor='black')
                ax.set_title('Проекты по статусам', fontsize=14, fontweight='bold')
                ax.set_ylabel('Количество', fontsize=12)
                ax.set_xlabel('Статус', fontsize=12)

                # Добавляем значения на столбцы
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.05,
                            f'{int(height)}', ha='center', va='bottom', fontsize=11, fontweight='bold')

                plt.xticks(rotation=0, fontsize=11)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, f'excel_status_{timestamp}.png')
                plt.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close()

                # Вставляем в Excel
                try:
                    excel_img = XLImage(chart_path)
                    excel_img.width = 400
                    excel_img.height = 250
                    worksheet.add_image(excel_img, 'D2')
                except Exception as e:
                    print(f"Не удалось вставить изображение в Excel: {e}")

        except Exception as e:
            print(f"Ошибка при создании изображений для Excel: {e}")

    def create_charts_for_reports(self, stats):
        """Создание графиков для отчетов"""
        try:
            charts_dir = 'reports/charts'
            if not os.path.exists(charts_dir):
                os.makedirs(charts_dir, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # График проектов по статусам
            if stats['projects_by_status']:
                fig, ax = plt.subplots(figsize=(10, 6))
                labels = list(stats['projects_by_status'].keys())
                values = list(stats['projects_by_status'].values())

                colors = ['#2E7D32', '#1565C0', '#EF6C00', '#C62828', '#6A1B9A']
                bars = ax.bar(labels, values, color=colors[:len(labels)],
                              edgecolor='black', linewidth=1.5)

                ax.set_title('Распределение проектов по статусам',
                             fontsize=16, fontweight='bold', pad=20)
                ax.set_ylabel('Количество проектов', fontsize=14)
                ax.set_xlabel('Статус проекта', fontsize=14)

                ax.grid(True, axis='y', linestyle='--', alpha=0.7)
                ax.set_axisbelow(True)

                # Добавляем значения на столбцы
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                            f'{int(height)}', ha='center', va='bottom',
                            fontsize=12, fontweight='bold')

                plt.xticks(fontsize=12)
                plt.yticks(fontsize=12)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, 'projects_by_status.png')
                plt.savefig(chart_path, dpi=200, bbox_inches='tight')
                plt.close()

            # График проектов по дисциплинам (если есть данные)
            if stats['projects_by_discipline']:
                fig, ax = plt.subplots(figsize=(12, 7))
                labels = list(stats['projects_by_discipline'].keys())
                values = list(stats['projects_by_discipline'].values())

                bars = ax.bar(labels, values, color='#2196F3',
                              edgecolor='black', linewidth=1.5)

                ax.set_title('Распределение проектов по дисциплинам',
                             fontsize=16, fontweight='bold', pad=20)
                ax.set_ylabel('Количество проектов', fontsize=14)
                ax.set_xlabel('Дисциплина', fontsize=14)

                ax.grid(True, axis='y', linestyle='--', alpha=0.7)
                ax.set_axisbelow(True)

                # Добавляем значения на столбцы
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                            f'{int(height)}', ha='center', va='bottom',
                            fontsize=11, fontweight='bold')

                plt.xticks(rotation=45, ha='right', fontsize=11)
                plt.yticks(fontsize=12)
                plt.tight_layout()

                chart_path = os.path.join(charts_dir, 'projects_by_discipline.png')
                plt.savefig(chart_path, dpi=200, bbox_inches='tight')
                plt.close()

        except Exception as e:
            print(f"Ошибка при создании графиков: {e}")

    def generate_excel_simple(self, stats):
        """Альтернативная генерация Excel без графиков"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Статистика"

            ws['A1'] = "Отчет по проектам"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws['A2'] = f"Сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
            ws['A2'].font = openpyxl.styles.Font(italic=True)

            row = 4
            ws[f'A{row}'] = "Общее количество проектов:"
            ws[f'B{row}'] = stats['total_projects']
            row += 1

            ws[f'A{row}'] = "Действий за 7 дней:"
            ws[f'B{row}'] = stats['actions_last_7_days']
            row += 1

            ws[f'A{row}'] = "Действий за 30 дней:"
            ws[f'B{row}'] = stats['actions_last_30_days']
            row += 2

            # Проекты по статусам
            ws[f'A{row}'] = "Проекты по статусам:"
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True)
            row += 1

            for status, count in stats['projects_by_status'].items():
                ws[f'A{row}'] = f"  {status}:"
                ws[f'B{row}'] = count
                row += 1

            row += 1

            # Проекты по дисциплинам
            ws[f'A{row}'] = "Проекты по дисциплинам:"
            ws[f'A{row}'].font = openpyxl.styles.Font(bold=True)
            row += 1

            for discipline, count in stats['projects_by_discipline'].items():
                ws[f'A{row}'] = f"  {discipline}:"
                ws[f'B{row}'] = count
                row += 1

            # Сохраняем файл
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir, exist_ok=True)

            excel_path = os.path.join(report_dir, f'projects_simple_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
            wb.save(excel_path)

            return os.path.abspath(excel_path)

        except Exception as e:
            raise Exception(f"Не удалось создать Excel отчет: {e}")

    def generate_word_report(self, stats):
        """Генерация Word отчета по требованиям"""
        try:
            # Создаем документ
            doc = Document()

            # Настройка стилей для делового документа
            self.setup_word_styles(doc)

            # Титульный лист
            self.add_title_page(doc)

            # Сводная таблица ключевых показателей
            self.add_summary_table(doc, stats)

            # Вставляем график
            self.add_charts_to_word(doc)

            # Последние проекты
            self.add_recent_projects(doc, stats)

            # Топ технологий
            self.add_top_technologies(doc, stats)

            # Сохраняем документ
            report_dir = 'reports'
            if not os.path.exists(report_dir):
                os.makedirs(report_dir, exist_ok=True)

            word_path = os.path.join(report_dir, 'projects_report.docx')
            doc.save(word_path)

            return os.path.abspath(word_path)

        except Exception as e:
            raise Exception(f"Не удалось создать Word отчет: {e}")

    def setup_word_styles(self, doc):
        """Настройка стилей Word документа"""
        # Основной стиль документа
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        # Стиль для заголовков
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            heading_style.font.size = Pt(16 - (i * 2))
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

        # Стиль для таблиц
        table_style = doc.styles.add_style('CustomTable', WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = 'Times New Roman'
        table_style.font.size = Pt(11)

    def add_title_page(self, doc):
        """Добавление титульного листа"""
        # Пустая страница для титульного листа
        section = doc.sections[0]

        # Заголовок
        title = doc.add_heading('ОТЧЕТ ПО ПРОЕКТАМ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пробелы
        for _ in range(5):
            doc.add_paragraph()

        # Информация
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y")}')
        p.runs[0].bold = True

        # Дополнительная информация
        for _ in range(10):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run('Сгенерировано системой управления проектами')
        p.runs[0].italic = True

        doc.add_page_break()

    def add_summary_table(self, doc, stats):
        """Добавление сводной таблицы"""
        doc.add_heading('Сводные показатели', level=1)

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = True

        # Настройка таблицы
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = 'CustomTable'

        # Данные
        data = [
            ("Общее количество проектов", str(stats['total_projects'])),
            ("Действий за 7 дней", str(stats['actions_last_7_days'])),
            ("Действий за 30 дней", str(stats['actions_last_30_days'])),
            ("Количество дисциплин", str(stats['disciplines_count'])),
            ("Уникальных технологий", str(stats.get('unique_technologies', 0))),
            ("Статусов проектов", str(len(stats['projects_by_status'])))
        ]

        for i, (label, value) in enumerate(data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value

            # Жирный шрифт для заголовков
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()

    def add_charts_to_word(self, doc):
        """Добавление графиков в Word"""
        doc.add_heading('Визуализация данных', level=1)

        charts_dir = 'reports/charts'

        # График по статусам
        status_chart_path = os.path.join(charts_dir, 'projects_by_status.png')
        if os.path.exists(status_chart_path):
            doc.add_heading('Распределение проектов по статусам', level=2)
            try:
                doc.add_picture(status_chart_path, width=Inches(6))
                # Центрируем изображение
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph("[График распределения по статусам]")

        doc.add_paragraph()

        # График по дисциплинам
        discipline_chart_path = os.path.join(charts_dir, 'projects_by_discipline.png')
        if os.path.exists(discipline_chart_path):
            doc.add_heading('Распределение проектов по дисциплинам', level=2)
            try:
                doc.add_picture(discipline_chart_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                doc.add_paragraph("[График распределения по дисциплинам]")

        doc.add_page_break()

    def add_recent_projects(self, doc, stats):
        """Добавление списка последних проектов"""
        doc.add_heading('Последние проекты', level=1)

        if stats['recent_projects']:
            table = doc.add_table(rows=len(stats['recent_projects']) + 1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.autofit = True

            # Заголовки
            headers = ["Название", "Дисциплина", "Статус", "Обновлен"]
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                # Жирный шрифт для заголовков
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Данные
            for i, project in enumerate(stats['recent_projects'], start=1):
                table.cell(i, 0).text = str(project[0]) if project[0] else ""
                table.cell(i, 1).text = str(project[1]) if project[1] else ""
                table.cell(i, 2).text = str(project[2]) if project[2] else ""

                if project[3]:
                    date_str = project[3].strftime('%d.%m.%Y %H:%M')
                    table.cell(i, 3).text = date_str

            # Настройка стиля таблицы
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = 'CustomTable'

        doc.add_paragraph()

    def add_top_technologies(self, doc, stats):
        """Добавление топ технологий"""
        if stats['top_technologies']:
            doc.add_heading('Топ используемых технологий', level=1)

            table = doc.add_table(rows=len(stats['top_technologies']) + 1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.autofit = True

            # Заголовки
            table.cell(0, 0).text = "Технология"
            table.cell(0, 1).text = "Количество использований"

            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Данные
            for i, (tech, count) in enumerate(stats['top_technologies'].items(), start=1):
                table.cell(i, 0).text = tech
                table.cell(i, 1).text = str(count)

            # Настройка стиля таблицы
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = 'CustomTable'

    def export_to_excel(self):
        """Экспорт текущего списка проектов в Excel"""
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Проекты"

            # Заголовок
            ws['A1'] = "Экспорт проектов"
            ws['A1'].font = openpyxl.styles.Font(bold=True, size=14)
            ws.merge_cells('A1:E1')
            ws['A1'].alignment = openpyxl.styles.Alignment(horizontal='center')

            ws['A2'] = f"Экспортировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            ws['A2'].font = openpyxl.styles.Font(italic=True)
            ws.merge_cells('A2:E2')
            ws['A2'].alignment = openpyxl.styles.Alignment(horizontal='center')

            # Заголовки колонок
            headers = ["Название", "Дисциплина", "Статус", "Дата создания", "Дата обновления"]
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=4, column=col, value=header)
                cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
                cell.fill = openpyxl.styles.PatternFill(start_color="366092",
                                                        end_color="366092",
                                                        fill_type="solid")

                # Настраиваем ширину колонок
                column_letter = openpyxl.utils.get_column_letter(col)
                ws.column_dimensions[column_letter].width = 20

            # Данные
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT name, discipline, status, created_at, updated_at 
                FROM projects 
                ORDER BY name
            """)

            for row_idx, project in enumerate(cursor.fetchall(), start=5):
                for col_idx, value in enumerate(project, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    if isinstance(value, datetime):
                        cell.value = value.strftime('%d.%m.%Y %H:%M')
                        cell.number_format = 'DD.MM.YYYY HH:MM'
                    else:
                        cell.value = value

                    # Чередование цветов строк
                    if row_idx % 2 == 0:
                        cell.fill = openpyxl.styles.PatternFill(start_color="F2F2F2",
                                                                end_color="F2F2F2",
                                                                fill_type="solid")

            cursor.close()
            conn.close()

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                initialfile=f"projects_export_{datetime.now().strftime('%Y%m%d')}.xlsx"
            )

            if file_path:
                wb.save(file_path)
                self.status_bar.config(text=f"Данные экспортированы в: {os.path.basename(file_path)}")
                messagebox.showinfo("Успех", f"Данные экспортированы в:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка экспорта: {str(e)}")

    def export_to_word(self):
        """Экспорт текущего проекта в Word"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для экспорта")
            return

        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cursor = conn.cursor()

            cursor.execute("""
                SELECT name, discipline, status, created_at, updated_at 
                FROM projects 
                WHERE id = %s
            """, (self.current_project_id,))

            project = cursor.fetchone()

            # Получаем технологии
            cursor.execute("""
                SELECT technology FROM technologies 
                WHERE project_id = %s 
                ORDER BY added_at
            """, (self.current_project_id,))

            technologies = [row[0] for row in cursor.fetchall()]

            cursor.close()
            conn.close()

            # Читаем описание из файла
            content = ""
            if self.current_project_file and os.path.exists(self.current_project_file):
                with open(self.current_project_file, 'r', encoding='utf-8') as f:
                    content = f.read()

            # Создаем Word документ
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.5

            # Заголовок
            title = doc.add_heading(project[0], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Информация о проекте
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_table.cell(0, 0).text = "Дисциплина:"
            info_table.cell(0, 1).text = project[1] or 'Не указана'

            info_table.cell(1, 0).text = "Статус:"
            info_table.cell(1, 1).text = project[2] or 'Не указан'

            if project[3]:
                created_date = project[3].strftime('%d.%m.%Y %H:%M')
                info_table.cell(2, 0).text = "Создан:"
                info_table.cell(2, 1).text = created_date

            if project[4]:
                updated_date = project[4].strftime('%d.%m.%Y %H:%M')
                info_table.cell(3, 0).text = "Обновлен:"
                info_table.cell(3, 1).text = updated_date

            doc.add_paragraph()

            # Технологии
            if technologies:
                doc.add_heading('Используемые технологии', level=2)
                for tech in technologies:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(tech)

            doc.add_heading('Описание проекта', level=2)

            # Конвертируем Markdown в структурированный текст
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(line[3:])
                elif line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()

            # Сохраняем файл
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                initialfile=f"{project[0].replace(' ', '_')}_report.docx"
            )

            if file_path:
                doc.save(file_path)
                self.status_bar.config(text=f"Проект экспортирован в: {os.path.basename(file_path)}")
                messagebox.showinfo("Успех", f"Проект экспортирован в:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать проект:\n{str(e)}")
            self.status_bar.config(text=f"Ошибка экспорта: {str(e)}")

    def show_quick_stats(self):
        """Показать быструю статистику"""
        try:
            stats = self.collect_statistics()

            stats_text = f"""
📊 БЫСТРАЯ СТАТИСТИКА
━━━━━━━━━━━━━━━━━━━━━━
• Всего проектов: {stats['total_projects']}
• Дисциплин: {stats['disciplines_count']}
• Действий за 7 дней: {stats['actions_last_7_days']}
• Действий за 30 дней: {stats['actions_last_30_days']}
• Уникальных технологий: {stats.get('unique_technologies', 0)}
━━━━━━━━━━━━━━━━━━━━━━
Статусы проектов:
"""
            for status, count in stats['projects_by_status'].items():
                stats_text += f"  • {status}: {count}\n"

            # Создаем окно со статистикой
            stats_window = tk.Toplevel(self.root)
            stats_window.title("Быстрая статистика")
            stats_window.geometry("400x400")

            text_widget = tk.Text(stats_window, wrap=tk.WORD, font=("Consolas", 10))
            text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            text_widget.insert(1.0, stats_text)
            text_widget.config(state=tk.DISABLED)

            # Кнопка закрытия
            ttk.Button(stats_window, text="Закрыть",
                       command=stats_window.destroy).pack(pady=5)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось получить статистику:\n{str(e)}")

    def on_closing(self):
        """Обработка закрытия приложения"""
        if messagebox.askokcancel("Выход", "Вы уверены, что хотите выйти?"):
            self.status_bar.config(text="Завершение работы...")
            self.root.after(100, self.root.destroy)


def check_dependencies():
    """Проверка необходимых библиотек"""
    required_libraries = {
        'psycopg2': 'psycopg2-binary',
        'openpyxl': 'openpyxl',
        'docx': 'python-docx',
        'matplotlib': 'matplotlib',
        'PIL': 'pillow',
        'markdown': 'markdown'
    }

    missing_libs = []
    for lib, pip_name in required_libraries.items():
        try:
            __import__(lib)
        except ImportError:
            missing_libs.append(pip_name)

    if missing_libs:
        print("❌ Ошибка: Не установлены необходимые библиотеки:")
        for lib in missing_libs:
            print(f"  - {lib}")
        print("\n📦 Установите зависимости командой:")
        print(f"pip install {' '.join(missing_libs)}")
        input("\nНажмите Enter для выхода...")
        return False
    return True


# Запуск приложения
if __name__ == "__main__":
    if check_dependencies():
        root = tk.Tk()
        app = ProjectManagerApp(root)
        root.mainloop()