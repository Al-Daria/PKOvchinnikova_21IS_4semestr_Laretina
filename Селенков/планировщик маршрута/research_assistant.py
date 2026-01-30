import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EducationalRoutePlanner:
    def __init__(self, root):
        self.root = root
        self.root.title("Планировщик индивидуального образовательного маршрута")
        self.root.geometry("1200x700")

        # Иконка окна
        try:
            self.root.iconbitmap('icon.ico')
        except:
            pass

        # Центрирование окна
        self.center_window()

        # Настройка цветовой схемы
        self.setup_colors()

        # Подключение к базе данных
        self.setup_database()

        # Загрузка данных
        self.load_competencies()

        # Загружаем достижения
        self.load_achievements_initial()

        # Статус бар
        self.status_bar = tk.Label(self.root,
                                   text="Готово к работе",
                                   bd=1,
                                   relief=tk.SUNKEN,
                                   anchor=tk.W,
                                   bg=self.colors['bg_dark'],
                                   fg=self.colors['text_light'])
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        # Основной интерфейс
        self.setup_ui()

        # Загрузка данных в интерфейс
        self.refresh_data()

        # Проверяем достижения
        self.check_achievements()

    def center_window(self):
        """Центрирование окна на экране"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def setup_colors(self):
        """Настройка цветовой схемы приложения"""
        self.colors = {
            'bg_light': '#f8fafc',
            'bg_dark': '#eef2f7',
            'primary': '#2c6fbb',
            'secondary': '#3b82f6',
            'success': '#10b981',
            'warning': '#f59e0b',
            'danger': '#ef4444',
            'info': '#06b6d4',
            'text': '#1e293b',
            'text_light': '#64748b',
            'border': '#cbd5e1',
            'highlight': '#e0f2fe',
            'tab_bg': '#ffffff',
            'tab_fg': '#475569',
            'tab_selected_bg': '#2c6fbb',
            'tab_selected_fg': '#ffffff',
            'card_bg': '#ffffff'
        }

    def setup_database(self):
        """Настройка подключения к базе данных"""
        try:
            import psycopg2
            self.conn = psycopg2.connect(
                user="postgres",
                password="1111",
                host="localhost",
                port="5432",
                database="postgres"
            )
            self.db_type = "postgres"
        except:
            self.conn = sqlite3.connect('educational_route.db', check_same_thread=False)
            self.db_type = "sqlite"

        self.cursor = self.conn.cursor()
        self.create_tables()

    def create_tables(self):
        """Создание таблиц в базе данных с русскими названиями как в ТЗ"""
        tables = [
            """CREATE TABLE IF NOT EXISTS цели (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                тип TEXT NOT NULL,
                статус TEXT NOT NULL,
                план_дата TEXT,
                факт_дата TEXT,
                описание TEXT
            )""",
            """CREATE TABLE IF NOT EXISTS навыки (
                id SERIAL PRIMARY KEY,
                название TEXT UNIQUE NOT NULL
            )""",
            """CREATE TABLE IF NOT EXISTS цель_навыки (
                id SERIAL PRIMARY KEY,
                цель_id INTEGER,
                навык_id INTEGER
            )""",
            """CREATE TABLE IF NOT EXISTS компетенции (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                категория TEXT
            )""",
            """CREATE TABLE IF NOT EXISTS цель_компетенции (
                id SERIAL PRIMARY KEY,
                цель_id INTEGER,
                компетенция_id INTEGER,
                уровень INTEGER CHECK (уровень >= 0 AND уровень <= 5)
            )""",
            """CREATE TABLE IF NOT EXISTS достижения (
                код TEXT PRIMARY KEY,
                название TEXT NOT NULL,
                описание TEXT,
                получено INTEGER DEFAULT 0
            )""",
            """CREATE TABLE IF NOT EXISTS цели_на_семестр (
                id SERIAL PRIMARY KEY,
                текст_цели TEXT NOT NULL,
                тип_цели TEXT,
                параметр TEXT,
                текущий_прогресс INTEGER DEFAULT 0,
                целевой_прогресс INTEGER NOT NULL
            )"""
        ]

        for table_sql in tables:
            try:
                self.cursor.execute(table_sql)
            except Exception as e:
                print(f"Ошибка при создании таблицы: {e}")

        self.conn.commit()

    def load_competencies(self):
        """Загрузка компетенций из файла"""
        try:
            if os.path.exists('competencies.json'):
                with open('competencies.json', 'r', encoding='utf-8') as f:
                    competencies = json.load(f)

                self.cursor.execute("SELECT COUNT(*) FROM компетенции")
                if self.cursor.fetchone()[0] == 0:
                    for comp in competencies:
                        if 'название' in comp and 'категория' in comp:
                            self.cursor.execute(
                                "INSERT INTO компетенции (название, категория) VALUES (%s, %s)"
                                if self.db_type == "postgres" else
                                "INSERT INTO компетенции (название, категория) VALUES (?, ?)",
                                (comp['название'], comp['категория'])
                            )
                    self.conn.commit()
        except Exception as e:
            print(f"Ошибка загрузки компетенций: {e}")

    def load_achievements_initial(self):
        """Первоначальная загрузка достижений без проверки"""
        achievements = [
            ("старт", "Старт", "Создана хотя бы одна цель"),
            ("пунктуальный", "Пунктуальный", "Три или более завершённых цели с фактической датой не позже плановой"),
            ("многоцелевой", "Многоцелевой", "Есть цели минимум трёх разных типов"),
            ("навыковый_рост", "Навыковый рост", "У одного навыка четыре или более связанных завершённых целей"),
            ("планировщик", "Планировщик", "Одновременно в статусе 'В процессе' пять или более целей")
        ]

        for ach in achievements:
            try:
                self.cursor.execute(
                    "INSERT INTO достижения (код, название, описание, получено) VALUES (%s, %s, %s, 0) ON CONFLICT (код) DO NOTHING"
                    if self.db_type == "postgres" else
                    "INSERT OR IGNORE INTO достижения (код, название, описание, получено) VALUES (?, ?, ?, 0)",
                    ach
                )
            except Exception as e:
                print(f"Ошибка загрузки достижений: {e}")

        self.conn.commit()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        self.setup_styles()

        # Создаем верхнюю панель с логотипом
        header_frame = tk.Frame(self.root, bg=self.colors['primary'], height=60)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)

        logo_label = tk.Label(header_frame,
                              text="🎓 Планировщик ИОМ",
                              font=('Arial', 16, 'bold'),
                              bg=self.colors['primary'],
                              fg='white')
        logo_label.pack(side=tk.LEFT, padx=20)

        # Блок с текущей датой
        date_label = tk.Label(header_frame,
                              text=datetime.now().strftime("%d.%m.%Y"),
                              font=('Arial', 10),
                              bg=self.colors['primary'],
                              fg='white')
        date_label.pack(side=tk.RIGHT, padx=20)

        self.notebook = ttk.Notebook(self.root, style="Custom.TNotebook")
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        # Создание вкладок с иконками
        self.goals_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.goals_frame, text="🎯 Мои цели")
        self.setup_goals_tab()

        self.profile_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.profile_frame, text="👤 Мой профиль")
        self.setup_profile_tab()

        self.competencies_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.competencies_frame, text="📊 Компетенции")
        self.setup_competencies_tab()

        self.achievements_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.achievements_frame, text="🏆 Достижения")
        self.setup_achievements_tab()

        self.semester_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.semester_frame, text="📅 Цели на семестр")
        self.setup_semester_tab()

        self.settings_frame = ttk.Frame(self.notebook, style="Custom.TFrame")
        self.notebook.add(self.settings_frame, text="⚙️ Настройки")
        self.setup_settings_tab()

    def setup_styles(self):
        """Настройка стилей виджетов"""
        style = ttk.Style()

        # Настраиваем тему
        style.theme_use('clam')

        # Стиль для Notebook
        style.configure("Custom.TNotebook",
                        background=self.colors['bg_light'],
                        tabmargins=[2, 5, 2, 0])

        style.configure("Custom.TNotebook.Tab",
                        background=self.colors['tab_bg'],
                        foreground=self.colors['tab_fg'],
                        padding=[20, 8],
                        font=('Arial', 10, 'bold'),
                        borderwidth=2,
                        relief=tk.FLAT)

        style.map("Custom.TNotebook.Tab",
                  background=[('selected', self.colors['tab_selected_bg']),
                              ('!selected', self.colors['tab_bg'])],
                  foreground=[('selected', self.colors['tab_selected_fg']),
                              ('!selected', self.colors['tab_fg'])],
                  relief=[('selected', 'raised')])

        # Стиль для фреймов
        style.configure("Custom.TFrame",
                        background=self.colors['bg_light'])

        # Стиль для Labelframe
        style.configure("Custom.TLabelframe",
                        background=self.colors['bg_light'],
                        foreground=self.colors['primary'],
                        bordercolor=self.colors['border'],
                        font=('Arial', 11, 'bold'),
                        relief=tk.GROOVE,
                        borderwidth=2)

        style.configure("Custom.TLabelframe.Label",
                        background=self.colors['bg_light'],
                        foreground=self.colors['primary'])

        # Стиль для кнопок
        style.configure("Primary.TButton",
                        background=self.colors['primary'],
                        foreground='white',
                        borderwidth=1,
                        focusthickness=3,
                        focuscolor='none',
                        font=('Arial', 10, 'bold'),
                        padding=10,
                        relief=tk.RAISED)

        style.map("Primary.TButton",
                  background=[('active', self.colors['secondary']),
                              ('pressed', '#1e5aa8'),
                              ('disabled', '#94a3b8')],
                  foreground=[('active', 'white'),
                              ('pressed', 'white'),
                              ('disabled', '#cbd5e1')])

        style.configure("Secondary.TButton",
                        background='white',
                        foreground=self.colors['primary'],
                        borderwidth=1,
                        border=self.colors['primary'],
                        focusthickness=3,
                        focuscolor='none',
                        font=('Arial', 10),
                        padding=8,
                        relief=tk.RAISED)

        style.map("Secondary.TButton",
                  background=[('active', self.colors['highlight']),
                              ('pressed', self.colors['primary'])],
                  foreground=[('active', self.colors['primary']),
                              ('pressed', 'white')])

        # Стиль для полей ввода
        style.configure("Custom.TEntry",
                        fieldbackground='white',
                        foreground=self.colors['text'],
                        bordercolor=self.colors['border'],
                        lightcolor=self.colors['border'],
                        darkcolor=self.colors['border'])

        style.map("Custom.TEntry",
                  fieldbackground=[('disabled', self.colors['bg_dark'])])

        # Стиль для Treeview
        style.configure("Custom.Treeview",
                        background='white',
                        foreground=self.colors['text'],
                        fieldbackground='white',
                        rowheight=28,
                        bordercolor=self.colors['border'],
                        borderwidth=1,
                        font=('Arial', 9))

        style.configure("Custom.Treeview.Heading",
                        background=self.colors['primary'],
                        foreground='white',
                        font=('Arial', 10, 'bold'),
                        relief=tk.FLAT)

        style.map("Custom.Treeview",
                  background=[('selected', self.colors['highlight'])],
                  foreground=[('selected', self.colors['text'])])

        # Стиль для Scrollbar
        style.configure("Custom.Vertical.TScrollbar",
                        background=self.colors['border'],
                        troughcolor=self.colors['bg_light'],
                        bordercolor=self.colors['border'],
                        arrowcolor=self.colors['primary'])

    def setup_goals_tab(self):
        """Настройка вкладки Мои цели - УПРОЩЕННАЯ ВЕРСИЯ"""
        # Главный контейнер с двумя панелями
        main_frame = tk.PanedWindow(self.goals_frame, orient=tk.HORIZONTAL, sashwidth=5, sashrelief=tk.RAISED)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Левая панель - форма (50%)
        left_frame = ttk.LabelFrame(main_frame, text="Добавить/редактировать цель", style="Custom.TLabelframe")
        main_frame.add(left_frame, width=600, minsize=300)

        form_frame = ttk.Frame(left_frame, style="Custom.TFrame")
        form_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Располагаем элементы в форме
        row = 0
        ttk.Label(form_frame, text="Название цели*:",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=(5, 3), padx=5)
        self.goal_name = ttk.Entry(form_frame, width=35, style="Custom.TEntry")
        self.goal_name.grid(row=row, column=1, pady=(5, 3), padx=5, sticky=tk.W)
        row += 1

        ttk.Label(form_frame, text="Тип*:",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=3, padx=5)
        self.goal_type = ttk.Combobox(form_frame,
                                      values=["Курс", "Проект", "Экзамен", "Исследование", "Практика", "Другое"],
                                      width=32, style="Custom.TCombobox")
        self.goal_type.grid(row=row, column=1, pady=3, padx=5, sticky=tk.W)
        row += 1

        ttk.Label(form_frame, text="Статус*:",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=3, padx=5)
        self.goal_status = ttk.Combobox(form_frame,
                                        values=["Планируется", "В процессе", "Завершено", "Отменено"],
                                        width=32, style="Custom.TCombobox")
        self.goal_status.grid(row=row, column=1, pady=3, padx=5, sticky=tk.W)
        row += 1

        ttk.Label(form_frame, text="Плановая дата (ГГГГ-ММ-ДД):",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=3, padx=5)
        self.goal_plan_date = ttk.Entry(form_frame, width=35, style="Custom.TEntry")
        self.goal_plan_date.grid(row=row, column=1, pady=3, padx=5, sticky=tk.W)
        self.goal_plan_date.insert(0, datetime.now().strftime("%Y-%m-%d"))
        row += 1

        ttk.Label(form_frame, text="Фактическая дата (ГГГГ-ММ-ДД):",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=3, padx=5)
        self.goal_fact_date = ttk.Entry(form_frame, width=35, style="Custom.TEntry")
        self.goal_fact_date.grid(row=row, column=1, pady=3, padx=5, sticky=tk.W)
        row += 1

        # Навыки - УПРОЩЕНО: просто текстовые поля
        ttk.Label(form_frame, text="Навыки (до 3х):",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=(10, 3), padx=5)
        row += 1

        self.skill_entries = []
        skill_frame = ttk.Frame(form_frame, style="Custom.TFrame")
        skill_frame.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=3, padx=5)

        for i in range(3):
            skill_entry = ttk.Entry(skill_frame, width=35, style="Custom.TEntry")
            skill_entry.grid(row=0, column=i, padx=(0, 5) if i < 2 else 0)
            self.skill_entries.append(skill_entry)
        row += 1

        # Компетенции - с фиксированным списком из базы
        ttk.Label(form_frame, text="Компетенции и уровни (до 3х):",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.W, pady=(10, 3), padx=5)
        row += 1

        # Получаем список компетенций из базы
        self.cursor.execute("SELECT название FROM компетенции ORDER BY название")
        competencies_list = [row[0] for row in self.cursor.fetchall()]

        self.competency_vars = []
        self.level_vars = []
        comp_frame = ttk.Frame(form_frame, style="Custom.TFrame")
        comp_frame.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=3, padx=5)

        for i in range(3):
            row_frame = ttk.Frame(comp_frame, style="Custom.TFrame")
            row_frame.pack(fill=tk.X, pady=2)

            competency_var = tk.StringVar()
            level_var = tk.StringVar(value="0")

            comp_combo = ttk.Combobox(row_frame, textvariable=competency_var,
                                      values=competencies_list, width=25, style="Custom.TCombobox")
            comp_combo.pack(side=tk.LEFT, padx=(0, 5))

            ttk.Label(row_frame, text="Уровень:", font=('Arial', 9)).pack(side=tk.LEFT, padx=(5, 2))

            level_combo = ttk.Combobox(row_frame, textvariable=level_var,
                                       values=["0", "1", "2", "3", "4", "5"],
                                       width=5, style="Custom.TCombobox")
            level_combo.pack(side=tk.LEFT)

            self.competency_vars.append(competency_var)
            self.level_vars.append(level_var)
        row += 1

        # Описание цели
        ttk.Label(form_frame, text="Описание:",
                  font=('Arial', 10, 'bold')).grid(row=row, column=0, sticky=tk.NW, pady=(10, 3), padx=5)
        row += 1

        # Контейнер для описания
        desc_container = ttk.Frame(form_frame, style="Custom.TFrame")
        desc_container.grid(row=row, column=0, columnspan=2, sticky=tk.W + tk.E, pady=3, padx=5)

        # Текстовое поле для ввода описания
        self.description_text = tk.Text(desc_container, height=5, width=40,
                                        bg='white', fg=self.colors['text'],
                                        font=('Arial', 9), wrap=tk.WORD)
        self.description_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Предпросмотр
        preview_frame = ttk.LabelFrame(form_frame, text="Предпросмотр", style="Custom.TLabelframe")
        preview_frame.grid(row=row + 1, column=0, columnspan=2, sticky=tk.W + tk.E, pady=10, padx=5)

        self.preview_text = tk.Text(preview_frame, height=4, state=tk.DISABLED,
                                    bg=self.colors['bg_dark'], fg=self.colors['text'],
                                    font=('Arial', 9), wrap=tk.WORD)
        self.preview_text.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

        # Кнопки формы
        button_frame = ttk.Frame(form_frame, style="Custom.TFrame")
        button_frame.grid(row=row + 2, column=0, columnspan=2, pady=10)

        ttk.Button(button_frame, text="💾 Сохранить цель",
                   command=self.save_goal,
                   style="Primary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(button_frame, text="👁️ Обновить предпросмотр",
                   command=self.update_preview,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(button_frame, text="🗑️ Очистить форму",
                   command=self.clear_form,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        # Правая панель - список целей (50%)
        right_frame = ttk.LabelFrame(main_frame, text="Список целей", style="Custom.TLabelframe")
        main_frame.add(right_frame, width=600, minsize=300)

        # Панель инструментов для списка
        list_toolbar = ttk.Frame(right_frame, style="Custom.TFrame")
        list_toolbar.pack(fill=tk.X, padx=5, pady=5)

        self.total_goals_count_label = ttk.Label(list_toolbar, text="Всего целей: 0",
                                                 font=('Arial', 10),
                                                 foreground=self.colors['primary'])
        self.total_goals_count_label.pack(side=tk.LEFT, padx=5)

        self.total_goals_label = ttk.Label(list_toolbar, text="",
                                           font=('Arial', 10, 'bold'),
                                           foreground=self.colors['success'])
        self.total_goals_label.pack(side=tk.LEFT, padx=5)

        # Treeview с прокруткой
        tree_frame = ttk.Frame(right_frame, style="Custom.TFrame")
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        columns = ("ID", "Название", "Тип", "Статус", "Срок")
        self.goals_tree = ttk.Treeview(tree_frame, columns=columns, show="headings",
                                       style="Custom.Treeview", selectmode="browse")

        for col in columns:
            self.goals_tree.heading(col, text=col, anchor=tk.CENTER)

        self.goals_tree.column("ID", width=50, anchor=tk.CENTER)
        self.goals_tree.column("Название", width=200, anchor=tk.W)
        self.goals_tree.column("Тип", width=100, anchor=tk.CENTER)
        self.goals_tree.column("Статус", width=100, anchor=tk.CENTER)
        self.goals_tree.column("Срок", width=100, anchor=tk.CENTER)

        self.goals_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.goals_tree.yview,
                                  style="Custom.Vertical.TScrollbar")
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.goals_tree.configure(yscrollcommand=scrollbar.set)

        # Кнопки управления списком
        list_buttons = ttk.Frame(right_frame, style="Custom.TFrame")
        list_buttons.pack(side=tk.BOTTOM, pady=10)

        ttk.Button(list_buttons, text="📥 Загрузить выбранную",
                   command=self.load_selected_goal,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(list_buttons, text="🗑️ Удалить выбранную",
                   command=self.delete_goal,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(list_buttons, text="🔄 Обновить список",
                   command=self.refresh_data,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        # Привязка событий
        self.description_text.bind('<KeyRelease>', lambda e: self.update_preview())
        self.goals_tree.bind('<<TreeviewSelect>>', lambda e: self.on_goal_tree_select())

    def update_preview(self):
        """Обновление предпросмотра разметки"""
        text = self.description_text.get("1.0", tk.END)
        self.preview_text.config(state=tk.NORMAL)
        self.preview_text.delete("1.0", tk.END)

        lines = text.split('\n')
        for line in lines:
            if line.startswith('- '):
                self.preview_text.insert(tk.END, '• ' + line[2:] + '\n')
            elif line.startswith('**') and line.endswith('**'):
                self.preview_text.insert(tk.END, line[2:-2].upper() + '\n')
            elif line.startswith('# '):
                self.preview_text.insert(tk.END, '📌 ' + line[2:] + '\n')
            elif '[' in line and '](' in line and ')' in line:
                start = line.find('[') + 1
                end = line.find(']')
                link_text = line[start:end]
                self.preview_text.insert(tk.END, f'🔗 {link_text}\n')
            else:
                self.preview_text.insert(tk.END, line + '\n')

        self.preview_text.config(state=tk.DISABLED)

    def on_goal_tree_select(self):
        """Обработка выбора цели в Treeview"""
        selection = self.goals_tree.selection()
        if selection:
            self.load_selected_goal_tree(selection[0])

    def load_selected_goal_tree(self, item_id):
        """Загрузка выбранной цели из Treeview"""
        values = self.goals_tree.item(item_id)['values']
        if values:
            goal_id = values[0]
            self.load_goal_by_id(goal_id)

    def load_goal_by_id(self, goal_id):
        """Загрузка цели по ID"""
        self.cursor.execute(
            "SELECT название, тип, статус, план_дата, факт_дата, описание FROM цели WHERE id = %s" if self.db_type == "postgres" else
            "SELECT название, тип, статус, план_дата, факт_дата, описание FROM цели WHERE id = ?",
            (goal_id,)
        )
        goal_data = self.cursor.fetchone()

        if goal_data:
            self.clear_form()

            name, goal_type, status, plan_date, fact_date, description = goal_data

            self.goal_name.delete(0, tk.END)
            self.goal_name.insert(0, name)
            self.goal_type.set(goal_type)
            self.goal_status.set(status)
            if plan_date:
                self.goal_plan_date.delete(0, tk.END)
                self.goal_plan_date.insert(0, plan_date)
            if fact_date:
                self.goal_fact_date.delete(0, tk.END)
                self.goal_fact_date.insert(0, fact_date)
            if description:
                self.description_text.delete("1.0", tk.END)
                self.description_text.insert("1.0", description)

            # Навыки
            self.cursor.execute(
                """SELECT н.название FROM навыки н
                JOIN цель_навыки цн ON н.id = цн.навык_id
                WHERE цн.цель_id = %s""" if self.db_type == "postgres" else
                """SELECT н.название FROM навыки н
                JOIN цель_навыки цн ON н.id = цн.навык_id
                WHERE цн.цель_id = ?""",
                (goal_id,)
            )
            skills = self.cursor.fetchall()

            for i, entry in enumerate(self.skill_entries):
                entry.delete(0, tk.END)

            for i, skill in enumerate(skills[:3]):
                if i < len(self.skill_entries):
                    self.skill_entries[i].delete(0, tk.END)
                    self.skill_entries[i].insert(0, skill[0])

            # Компетенции
            self.cursor.execute(
                """SELECT к.название, цк.уровень FROM компетенции к
                JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                WHERE цк.цель_id = %s""" if self.db_type == "postgres" else
                """SELECT к.название, цк.уровень FROM компетенции к
                JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                WHERE цк.цель_id = ?""",
                (goal_id,)
            )
            comps = self.cursor.fetchall()

            for var in self.competency_vars:
                var.set('')
            for var in self.level_vars:
                var.set('0')

            for i, comp in enumerate(comps[:3]):
                if i < len(self.competency_vars):
                    self.competency_vars[i].set(comp[0])
                    self.level_vars[i].set(str(comp[1]))

            self.update_preview()

    def save_goal(self):
        """Сохранение цели в базу данных"""
        name = self.goal_name.get()
        goal_type = self.goal_type.get()
        status = self.goal_status.get()
        plan_date = self.goal_plan_date.get()
        fact_date = self.goal_fact_date.get()
        description = self.description_text.get("1.0", tk.END).strip()

        if not name or not goal_type or not status:
            messagebox.showerror("Ошибка", "Заполните обязательные поля: название, тип и статус")
            return

        try:
            if plan_date:
                datetime.strptime(plan_date, '%Y-%m-%d')
            if fact_date:
                datetime.strptime(fact_date, '%Y-%m-%d')
        except ValueError:
            messagebox.showerror("Ошибка", "Неверный формат даты. Используйте ГГГГ-ММ-ДД")
            return

        try:
            self.cursor.execute(
                """INSERT INTO цели (название, тип, статус, план_дата, факт_дата, описание) 
                VALUES (%s, %s, %s, %s, %s, %s) RETURNING id"""
                if self.db_type == "postgres" else
                """INSERT INTO цели (название, тип, статус, план_дата, факт_дата, описание) 
                VALUES (?, ?, ?, ?, ?, ?)""",
                (name, goal_type, status, plan_date or None, fact_date or None, description)
            )

            if self.db_type == "postgres":
                goal_id = self.cursor.fetchone()[0]
            else:
                goal_id = self.cursor.lastrowid

            # Сохраняем навыки
            for skill_entry in self.skill_entries[:3]:
                skill = skill_entry.get().strip()
                if skill:
                    # Проверяем, есть ли такой навык в базе
                    self.cursor.execute(
                        "SELECT id FROM навыки WHERE название = %s" if self.db_type == "postgres" else
                        "SELECT id FROM навыки WHERE название = ?",
                        (skill,)
                    )
                    result = self.cursor.fetchone()

                    if result:
                        skill_id = result[0]
                    else:
                        self.cursor.execute(
                            "INSERT INTO навыки (название) VALUES (%s) RETURNING id" if self.db_type == "postgres" else
                            "INSERT INTO навыки (название) VALUES (?)",
                            (skill,)
                        )
                        if self.db_type == "postgres":
                            skill_id = self.cursor.fetchone()[0]
                        else:
                            skill_id = self.cursor.lastrowid

                    self.cursor.execute(
                        "INSERT INTO цель_навыки (цель_id, навык_id) VALUES (%s, %s)" if self.db_type == "postgres" else
                        "INSERT INTO цель_навыки (цель_id, навык_id) VALUES (?, ?)",
                        (goal_id, skill_id)
                    )

            # Сохраняем компетенции
            for i in range(3):
                comp_name = self.competency_vars[i].get().strip()
                level = self.level_vars[i].get().strip()

                if comp_name and level:
                    # Получаем ID компетенции
                    self.cursor.execute(
                        "SELECT id FROM компетенции WHERE название = %s" if self.db_type == "postgres" else
                        "SELECT id FROM компетенции WHERE название = ?",
                        (comp_name,)
                    )
                    result = self.cursor.fetchone()

                    if result:
                        comp_id = result[0]
                        self.cursor.execute(
                            "INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень) VALUES (%s, %s, %s)" if self.db_type == "postgres" else
                            "INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень) VALUES (?, ?, ?)",
                            (goal_id, comp_id, int(level))
                        )

            self.conn.commit()
            self.status_bar.config(text="Цель успешно сохранена!")
            messagebox.showinfo("Успех", "Цель сохранена!")
            self.refresh_data()
            self.clear_form()
            # Проверяем достижения после сохранения цели
            self.check_achievements()

        except Exception as e:
            self.status_bar.config(text=f"Ошибка при сохранении: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")

    def refresh_data(self):
        """Обновление данных во всех вкладках"""
        self.refresh_goals_list()
        self.refresh_profile()
        self.refresh_competencies()
        self.refresh_achievements()
        self.refresh_semester_goals()

    def refresh_goals_list(self):
        """Обновление списка целей"""
        for item in self.goals_tree.get_children():
            self.goals_tree.delete(item)

        self.cursor.execute("SELECT id, название, тип, статус, план_дата FROM цели ORDER BY план_дата DESC, id DESC")
        goals = self.cursor.fetchall()

        for goal in goals:
            goal_id, name, goal_type, status, plan_date = goal
            display_date = plan_date if plan_date else "—"

            self.goals_tree.insert("", tk.END, values=(goal_id, name, goal_type, status, display_date))

        # Обновление счетчика
        total_count = len(goals)
        completed_count = len([g for g in goals if g[3] == 'Завершено'])

        # Обновляем метку "Всего целей"
        self.total_goals_count_label.config(text=f"Всего целей: {total_count}")
        self.total_goals_label.config(
            text=f"({completed_count} завершено, {total_count - completed_count} в работе)"
        )

    def load_selected_goal(self):
        """Загрузка выбранной цели в форму"""
        selection = self.goals_tree.selection()
        if not selection:
            messagebox.showwarning("Внимание", "Выберите цель из списка")
            return

        self.load_selected_goal_tree(selection[0])

    def delete_goal(self):
        """Удаление выбранной цели"""
        selection = self.goals_tree.selection()
        if not selection:
            messagebox.showwarning("Внимание", "Выберите цель из списка")
            return

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            values = self.goals_tree.item(selection[0])['values']
            if values:
                goal_id = values[0]

                try:
                    self.cursor.execute(
                        "DELETE FROM цель_навыки WHERE цель_id = %s" if self.db_type == "postgres" else
                        "DELETE FROM цель_навыки WHERE цель_id = ?",
                        (goal_id,)
                    )
                    self.cursor.execute(
                        "DELETE FROM цель_компетенции WHERE цель_id = %s" if self.db_type == "postgres" else
                        "DELETE FROM цель_компетенции WHERE цель_id = ?",
                        (goal_id,)
                    )
                    self.cursor.execute(
                        "DELETE FROM цели WHERE id = %s" if self.db_type == "postgres" else
                        "DELETE FROM цели WHERE id = ?",
                        (goal_id,)
                    )

                    self.conn.commit()
                    self.refresh_data()
                    self.status_bar.config(text="Цель успешно удалена!")
                    messagebox.showinfo("Успех", "Цель удалена!")

                except Exception as e:
                    self.status_bar.config(text=f"Ошибка при удалении: {str(e)}")
                    messagebox.showerror("Ошибка", f"Ошибка при удалении: {str(e)}")

    def clear_form(self):
        """Очистка формы ввода"""
        self.goal_name.delete(0, tk.END)
        self.goal_type.set('')
        self.goal_status.set('')
        self.goal_plan_date.delete(0, tk.END)
        self.goal_plan_date.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.goal_fact_date.delete(0, tk.END)
        self.description_text.delete("1.0", tk.END)

        for entry in self.skill_entries:
            entry.delete(0, tk.END)

        for var in self.competency_vars:
            var.set('')

        for var in self.level_vars:
            var.set('0')

        self.update_preview()

    def setup_profile_tab(self):
        """Настройка вкладки Мой профиль"""
        main_container = ttk.Frame(self.profile_frame, style="Custom.TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Навыки
        skills_frame = ttk.LabelFrame(main_container, text="Навыки", style="Custom.TLabelframe")
        skills_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.skills_text = tk.Text(skills_frame, height=10, state=tk.DISABLED,
                                   bg=self.colors['bg_dark'], fg=self.colors['text'],
                                   font=('Arial', 9), wrap=tk.WORD)
        self.skills_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Статистика
        stats_frame = ttk.LabelFrame(main_container, text="Статистика", style="Custom.TLabelframe")
        stats_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.stats_text = tk.Text(stats_frame, height=10, state=tk.DISABLED,
                                  bg=self.colors['bg_dark'], fg=self.colors['text'],
                                  font=('Arial', 9), wrap=tk.WORD)
        self.stats_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    def refresh_profile(self):
        """Обновление вкладки профиля"""
        self.skills_text.config(state=tk.NORMAL)
        self.skills_text.delete("1.0", tk.END)

        self.cursor.execute("""
            SELECT н.название, COUNT(цн.цель_id) as количество_целей,
                   SUM(CASE WHEN ц.статус = 'Завершено' THEN 1 ELSE 0 END) as завершено_целей
            FROM навыки н
            LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
            LEFT JOIN цели ц ON цн.цель_id = ц.id
            GROUP BY н.id, н.название
            ORDER BY количество_целей DESC
        """)

        skills = self.cursor.fetchall()
        if skills:
            for skill, total, completed in skills:
                status_icon = "✅" if completed > 0 else "⏳"
                self.skills_text.insert(tk.END, f"{status_icon} {skill} — всего: {total}")
                if completed > 0:
                    self.skills_text.insert(tk.END, f" (завершено: {completed})")
                self.skills_text.insert(tk.END, "\n")
        else:
            self.skills_text.insert(tk.END, "Навыки не указаны\n")

        self.skills_text.config(state=tk.DISABLED)

        self.stats_text.config(state=tk.NORMAL)
        self.stats_text.delete("1.0", tk.END)

        self.cursor.execute("""
            SELECT тип, 
                   COUNT(*) as всего,
                   SUM(CASE WHEN статус = 'Завершено' THEN 1 ELSE 0 END) as завершено
            FROM цели
            GROUP BY тип
        """)

        type_stats = self.cursor.fetchall()
        for type_name, total, completed in type_stats:
            progress = (completed / total * 100) if total > 0 else 0
            progress_bar = self.get_progress_bar(progress)
            self.stats_text.insert(tk.END, f"📊 {type_name}: {completed} из {total}\n")
            self.stats_text.insert(tk.END, f"   {progress_bar} {progress:.0f}%\n\n")

        self.stats_text.insert(tk.END, "\n")

        self.cursor.execute("""
            SELECT 
                COUNT(*) as всего_завершённых,
                SUM(CASE WHEN факт_дата <= план_дата THEN 1 ELSE 0 END) as в_срок
            FROM цели
            WHERE статус = 'Завершено' AND план_дата IS NOT NULL AND факт_дата IS NOT NULL
        """)

        result = self.cursor.fetchone()
        if result and result[0] > 0:
            total_completed, on_time = result
            percentage = (on_time / total_completed) * 100 if total_completed > 0 else 0
            progress_bar = self.get_progress_bar(percentage)
            self.stats_text.insert(tk.END, f"⏰ Целей завершено в срок:\n")
            self.stats_text.insert(tk.END, f"   {on_time} из {total_completed}\n")
            self.stats_text.insert(tk.END, f"   {progress_bar} {percentage:.1f}%\n")
        else:
            self.stats_text.insert(tk.END, "⏰ Целей завершено в срок: нет данных\n")

        self.stats_text.config(state=tk.DISABLED)

    def get_progress_bar(self, percentage, length=20):
        """Создание текстового прогресс-бара"""
        filled = int((percentage / 100) * length)
        empty = length - filled
        return "█" * filled + "░" * empty

    def setup_competencies_tab(self):
        """Настройка вкладки Компетенции"""
        main_container = ttk.Frame(self.competencies_frame, style="Custom.TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        avg_frame = ttk.LabelFrame(main_container, text="Средний уровень по компетенциям",
                                   style="Custom.TLabelframe")
        avg_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.avg_text = tk.Text(avg_frame, height=10, state=tk.DISABLED,
                                bg=self.colors['bg_dark'], fg=self.colors['text'],
                                font=('Arial', 9), wrap=tk.WORD)
        self.avg_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        weak_frame = ttk.LabelFrame(main_container, text="⚠️ Слабые зоны", style="Custom.TLabelframe")
        weak_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.weak_text = tk.Text(weak_frame, height=5, state=tk.DISABLED,
                                 bg='#fff3cd', fg=self.colors['text'],
                                 font=('Arial', 9), wrap=tk.WORD)
        self.weak_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        rec_frame = ttk.LabelFrame(main_container, text="💡 Рекомендации", style="Custom.TLabelframe")
        rec_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.rec_text = tk.Text(rec_frame, height=8, state=tk.DISABLED,
                                bg='#d1ecf1', fg=self.colors['text'],
                                font=('Arial', 9), wrap=tk.WORD)
        self.rec_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    def refresh_competencies(self):
        """Обновление вкладки компетенций"""
        self.avg_text.config(state=tk.NORMAL)
        self.avg_text.delete("1.0", tk.END)

        self.cursor.execute("""
            SELECT к.название, к.категория, ROUND(AVG(цк.уровень), 1) as средний_уровень
            FROM компетенции к
            LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
            GROUP BY к.id, к.название, к.категория
            ORDER BY средний_уровень DESC NULLS LAST
        """)

        comps = self.cursor.fetchall()
        for name, category, avg_level in comps:
            if avg_level:
                level_int = int(float(avg_level))
                stars = "★" * level_int + "☆" * (5 - level_int)
                level_text = f"{avg_level:.1f}"

                self.avg_text.insert(tk.END, f"• {name} ({category}): {level_text} {stars}\n")
            else:
                self.avg_text.insert(tk.END, f"• {name} ({category}): нет данных\n")

        self.avg_text.config(state=tk.DISABLED)

        self.weak_text.config(state=tk.NORMAL)
        self.weak_text.delete("1.0", tk.END)

        weak_zones = []
        for name, category, avg_level in comps:
            if avg_level and avg_level < 3:
                weak_zones.append((name, category, avg_level))

        if weak_zones:
            self.weak_text.insert(tk.END, "⚠️ Внимание! Эти компетенции требуют развития:\n\n")
            for name, category, avg_level in weak_zones:
                self.weak_text.insert(tk.END, f"• {name} ({category}): уровень {avg_level:.1f}\n")
        else:
            self.weak_text.insert(tk.END, "✅ Все компетенции развиты достаточно хорошо\n")

        self.weak_text.config(state=tk.DISABLED)

        self.rec_text.config(state=tk.NORMAL)
        self.rec_text.delete("1.0", tk.END)

        if weak_zones:
            self.rec_text.insert(tk.END, "💡 Рекомендации по развитию:\n\n")
            for name, category, avg_level in weak_zones[:3]:
                if "презентация" in name.lower():
                    self.rec_text.insert(tk.END,
                                         f"• Для компетенции '{name}': выступите на студенческой конференции или подготовьте презентацию для семинара.\n\n")
                elif "баз" in name.lower() or "данн" in name.lower():
                    self.rec_text.insert(tk.END,
                                         f"• Для компетенции '{name}': пройдите курс по базам данных на Stepik или выполните учебный проект.\n\n")
                elif "проект" in name.lower():
                    self.rec_text.insert(tk.END,
                                         f"• Для компетенции '{name}': участвуйте в командных проектах или организуйте собственный мини-проект.\n\n")
                else:
                    self.rec_text.insert(tk.END,
                                         f"• Для компетенции '{name}': ищите практические задания и кейсы по этой теме.\n\n")
        else:
            self.rec_text.insert(tk.END, "🎉 Все компетенции развиты хорошо. Продолжайте в том же духе!\n")

        self.rec_text.config(state=tk.DISABLED)

    def setup_achievements_tab(self):
        """Настройка вкладки Достижения"""
        main_container = ttk.Frame(self.achievements_frame, style="Custom.TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        header_frame = ttk.Frame(main_container, style="Custom.TFrame")
        header_frame.pack(fill=tk.X, pady=(0, 20))

        header_label = ttk.Label(header_frame,
                                 text="🏆 Ваши достижения",
                                 font=('Arial', 14, 'bold'),
                                 foreground=self.colors['primary'],
                                 background=self.colors['bg_light'])
        header_label.pack()

        self.achievements_container = ttk.Frame(main_container, style="Custom.TFrame")
        self.achievements_container.pack(fill=tk.BOTH, expand=True)

    def refresh_achievements(self):
        """Обновление вкладки достижений"""
        for widget in self.achievements_container.winfo_children():
            widget.destroy()

        self.cursor.execute("SELECT код, название, описание, получено FROM достижения ORDER BY получено DESC, код")
        achievements = self.cursor.fetchall()

        if not achievements:
            no_ach_label = ttk.Label(self.achievements_container,
                                     text="Достижения пока не получены",
                                     font=('Arial', 11),
                                     foreground=self.colors['text_light'],
                                     background=self.colors['bg_light'])
            no_ach_label.pack(pady=50)
            return

        obtained = [a for a in achievements if a[3] == 1]
        not_obtained = [a for a in achievements if a[3] == 0]

        if obtained:
            ttk.Label(self.achievements_container,
                      text="✅ Полученные достижения:",
                      font=('Arial', 12, 'bold'),
                      foreground=self.colors['success'],
                      background=self.colors['bg_light']).pack(anchor=tk.W, pady=(0, 10))

            for code, name, description, obtained_status in obtained:
                self.create_achievement_card(self.achievements_container, name, description, True)

        if not_obtained:
            ttk.Label(self.achievements_container,
                      text="⏳ Достижения в процессе:",
                      font=('Arial', 12, 'bold'),
                      foreground=self.colors['warning'],
                      background=self.colors['bg_light']).pack(anchor=tk.W, pady=(20, 10))

            for code, name, description, obtained_status in not_obtained:
                self.create_achievement_card(self.achievements_container, name, description, False)

        total = len(achievements)
        obtained_count = len(obtained)
        progress = (obtained_count / total * 100) if total > 0 else 0

        stats_frame = ttk.Frame(self.achievements_container, style="Custom.TFrame")
        stats_frame.pack(fill=tk.X, pady=20)

        stats_text = f"📊 Прогресс: {obtained_count} из {total} достижений ({progress:.1f}%)"
        stats_label = ttk.Label(stats_frame,
                                text=stats_text,
                                font=('Arial', 10),
                                foreground=self.colors['primary'],
                                background=self.colors['bg_light'])
        stats_label.pack()

    def create_achievement_card(self, parent, name, description, obtained):
        """Создание карточки достижения"""
        card_frame = ttk.Frame(parent, style="Custom.TFrame")
        card_frame.pack(fill=tk.X, pady=5, padx=5)

        inner_frame = tk.Frame(card_frame, bg=self.colors['bg_dark'] if not obtained else '#d4edda',
                               relief=tk.RIDGE, borderwidth=1)
        inner_frame.pack(fill=tk.X, padx=2, pady=2)

        icon_label = tk.Label(inner_frame,
                              text="✅" if obtained else "⏳",
                              font=('Arial', 14),
                              bg=inner_frame['bg'])
        icon_label.pack(side=tk.LEFT, padx=10, pady=10)

        text_frame = tk.Frame(inner_frame, bg=inner_frame['bg'])
        text_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10), pady=10)

        name_label = tk.Label(text_frame,
                              text=name,
                              font=('Arial', 11, 'bold'),
                              bg=inner_frame['bg'],
                              fg=self.colors['primary'] if obtained else self.colors['text'])
        name_label.pack(anchor=tk.W)

        desc_label = tk.Label(text_frame,
                              text=description,
                              font=('Arial', 9),
                              bg=inner_frame['bg'],
                              fg=self.colors['text'],
                              wraplength=500,
                              justify=tk.LEFT)
        desc_label.pack(anchor=tk.W, pady=(2, 0))

        status_frame = tk.Frame(inner_frame, bg=inner_frame['bg'])
        status_frame.pack(side=tk.RIGHT, padx=10, pady=10)

        status_text = "Получено" if obtained else "В процессе"
        status_color = self.colors['success'] if obtained else self.colors['warning']

        status_label = tk.Label(status_frame,
                                text=status_text,
                                font=('Arial', 9, 'bold'),
                                bg=status_color,
                                fg='white',
                                padx=10,
                                pady=2)
        status_label.pack()

    def check_achievements(self):
        """Проверка и обновление достижений - ИСПРАВЛЕННАЯ ВЕРСИЯ"""
        try:
            self.status_bar.config(text="Проверка достижений...")

            # Проверка всех достижений из ТЗ
            checks = [
                ("старт", "SELECT COUNT(*) FROM цели"),
                ("пунктуальный", """SELECT COUNT(*) FROM цели 
                                   WHERE статус = 'Завершено' 
                                   AND факт_дата IS NOT NULL 
                                   AND план_дата IS NOT NULL 
                                   AND факт_дата <= план_дата"""),
                ("многоцелевой", "SELECT COUNT(DISTINCT тип) FROM цели"),
                ("навыковый_рост", """SELECT COUNT(*) FROM (
                    SELECT н.id 
                    FROM навыки н
                    JOIN цель_навыки цн ON н.id = цн.навык_id
                    JOIN цели ц ON цн.цель_id = ц.id
                    WHERE ц.статус = 'Завершено'
                    GROUP BY н.id
                    HAVING COUNT(ц.id) >= 4
                ) as skill_counts"""),
                ("планировщик", "SELECT COUNT(*) FROM цели WHERE статус = 'В процессе'")
            ]

            for code, query in checks:
                try:
                    self.cursor.execute(query)
                    result = self.cursor.fetchone()

                    if result:
                        count = result[0]
                        # Определяем условия для каждого достижения
                        if code == "старт":
                            achieved = count > 0
                        elif code == "пунктуальный":
                            achieved = count >= 3
                        elif code == "многоцелевой":
                            achieved = count >= 3
                        elif code == "навыковый_рост":
                            achieved = count > 0
                        elif code == "планировщик":
                            achieved = count >= 5
                        else:
                            achieved = False
                    else:
                        achieved = False

                    # Обновляем достижение
                    self.cursor.execute(
                        "UPDATE достижения SET получено = %s WHERE код = %s" if self.db_type == "postgres" else
                        "UPDATE достижения SET получено = ? WHERE код = ?",
                        (1 if achieved else 0, code)
                    )
                except Exception as e:
                    print(f"Ошибка при проверке достижения {code}: {e}")

            self.conn.commit()
            self.status_bar.config(text="Достижения проверены")
            self.refresh_achievements()

        except Exception as e:
            self.status_bar.config(text=f"Ошибка проверки достижений: {e}")

    def setup_semester_tab(self):
        """Настройка вкладки Цели на семестр"""
        main_container = ttk.Frame(self.semester_frame, style="Custom.TFrame")
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Форма добавления цели
        form_frame = ttk.LabelFrame(main_container, text="Добавить цель на семестр", style="Custom.TLabelframe")
        form_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(form_frame, text="Текст цели*:", font=('Arial', 10, 'bold')).grid(row=0, column=0, sticky=tk.W,
                                                                                    pady=5,
                                                                                    padx=5)
        self.semester_goal_text = ttk.Entry(form_frame, width=50, style="Custom.TEntry")
        self.semester_goal_text.grid(row=0, column=1, pady=5, padx=5, sticky=tk.W)

        ttk.Label(form_frame, text="Тип цели*:", font=('Arial', 10, 'bold')).grid(row=1, column=0, sticky=tk.W, pady=5,
                                                                                  padx=5)
        self.semester_goal_type = ttk.Combobox(form_frame, values=["Количество", "Повысить компетенцию"],
                                               width=47, style="Custom.TCombobox")
        self.semester_goal_type.grid(row=1, column=1, pady=5, padx=5, sticky=tk.W)

        ttk.Label(form_frame, text="Параметр:", font=('Arial', 10, 'bold')).grid(row=2, column=0, sticky=tk.W, pady=5,
                                                                                 padx=5)
        self.semester_goal_param = ttk.Entry(form_frame, width=50, style="Custom.TEntry")
        self.semester_goal_param.grid(row=2, column=1, pady=5, padx=5, sticky=tk.W)

        ttk.Label(form_frame, text="Целевой прогресс*:", font=('Arial', 10, 'bold')).grid(row=3, column=0, sticky=tk.W,
                                                                                          pady=5, padx=5)
        self.semester_target_progress = ttk.Spinbox(form_frame, from_=1, to=100, width=48, style="Custom.TEntry")
        self.semester_target_progress.grid(row=3, column=1, pady=5, padx=5, sticky=tk.W)
        self.semester_target_progress.delete(0, tk.END)
        self.semester_target_progress.insert(0, '1')

        ttk.Button(form_frame, text="➕ Добавить цель", command=self.add_semester_goal,
                   style="Primary.TButton").grid(row=4, column=0, columnspan=2, pady=10)

        # Список целей
        list_frame = ttk.LabelFrame(main_container, text="Цели на семестр", style="Custom.TLabelframe")
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        columns = ("ID", "Цель", "Прогресс", "Визуализация")
        self.semester_tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=8,
                                          style="Custom.Treeview")

        self.semester_tree.heading("ID", text="ID", anchor=tk.CENTER)
        self.semester_tree.heading("Цель", text="Цель", anchor=tk.W)
        self.semester_tree.heading("Прогресс", text="Прогресс", anchor=tk.CENTER)
        self.semester_tree.heading("Визуализация", text="Визуализация", anchor=tk.CENTER)

        self.semester_tree.column("ID", width=50, anchor=tk.CENTER)
        self.semester_tree.column("Цель", width=300, anchor=tk.W)
        self.semester_tree.column("Прогресс", width=120, anchor=tk.CENTER)
        self.semester_tree.column("Визуализация", width=150, anchor=tk.CENTER)

        self.semester_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.semester_tree.yview,
                                  style="Custom.Vertical.TScrollbar")
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.semester_tree.configure(yscrollcommand=scrollbar.set)

        # Кнопки управления
        button_frame = ttk.Frame(main_container, style="Custom.TFrame")
        button_frame.pack(pady=10)

        ttk.Button(button_frame, text="📈 Изменить прогресс", command=self.update_semester_progress,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(button_frame, text="🗑️ Удалить", command=self.delete_semester_goal,
                   style="Secondary.TButton").pack(side=tk.LEFT, padx=5)

        ttk.Button(main_container, text="📄 Сформировать отчёт", command=self.generate_report,
                   style="Primary.TButton").pack(pady=10)

        self.semester_tree.bind('<<TreeviewSelect>>', self.on_semester_goal_select)

    def on_semester_goal_select(self, event):
        """Обработка выбора цели в Treeview"""
        selection = self.semester_tree.selection()
        if selection:
            self.selected_semester_id = self.semester_tree.item(selection[0])['values'][0]
        else:
            self.selected_semester_id = None

    def add_semester_goal(self):
        """Добавление цели на семестр"""
        text = self.semester_goal_text.get()
        goal_type = self.semester_goal_type.get()
        param = self.semester_goal_param.get()
        target = self.semester_target_progress.get()

        if not text or not goal_type or not target:
            messagebox.showerror("Ошибка", "Заполните обязательные поля")
            return

        try:
            target_int = int(target)
            if target_int <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Ошибка", "Целевой прогресс должен быть положительным числом")
            return

        try:
            self.cursor.execute(
                """INSERT INTO цели_на_семестр (текст_цели, тип_цели, параметр, целевой_прогресс) 
                VALUES (%s, %s, %s, %s)""" if self.db_type == "postgres" else
                """INSERT INTO цели_на_семестр (текст_цели, тип_цели, параметр, целевой_прогресс) 
                VALUES (?, ?, ?, ?)""",
                (text, goal_type, param, target_int)
            )

            self.conn.commit()
            self.status_bar.config(text="Цель на семестр добавлена!")
            messagebox.showinfo("Успех", "Цель на семестр добавлена!")

            self.semester_goal_text.delete(0, tk.END)
            self.semester_goal_type.set('')
            self.semester_goal_param.delete(0, tk.END)
            self.semester_target_progress.delete(0, tk.END)
            self.semester_target_progress.insert(0, '1')

            self.refresh_semester_goals()

        except Exception as e:
            self.status_bar.config(text=f"Ошибка при добавлении: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при добавлении: {str(e)}")

    def update_semester_progress(self):
        """Изменение прогресса выполнения цели на семестр"""
        if not hasattr(self, 'selected_semester_id') or not self.selected_semester_id:
            messagebox.showwarning("Внимание", "Выберите цель из списка")
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Изменение прогресса")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.configure(bg=self.colors['bg_light'])

        self.cursor.execute(
            "SELECT текст_цели, текущий_прогресс, целевой_прогресс FROM цели_на_семестр WHERE id = %s" if self.db_type == "postgres" else
            "SELECT текст_цели, текущий_прогресс, целевой_прогресс FROM цели_на_семестр WHERE id = ?",
            (self.selected_semester_id,)
        )
        goal_data = self.cursor.fetchone()

        if goal_data:
            text, current, target = goal_data

            info_label = ttk.Label(dialog,
                                   text=f"Цель: {text}\nТекущий прогресс: {current} из {target}",
                                   background=self.colors['bg_light'],
                                   font=('Arial', 9))
            info_label.pack(pady=10)

        ttk.Label(dialog, text="Введите новый прогресс (0-100%):", background=self.colors['bg_light']).pack(pady=5)

        progress_var = tk.StringVar()
        progress_spinbox = ttk.Spinbox(dialog, from_=0, to=100, textvariable=progress_var, width=20,
                                       style="Custom.TEntry")
        progress_spinbox.pack(pady=10)

        def save_progress():
            try:
                progress = int(progress_var.get())
                if progress < 0 or progress > 100:
                    raise ValueError

                actual_progress = int((progress / 100) * target)

                self.cursor.execute(
                    "UPDATE цели_на_семестр SET текущий_прогресс = %s WHERE id = %s" if self.db_type == "postgres" else
                    "UPDATE цели_на_семестр SET текущий_прогресс = ? WHERE id = ?",
                    (actual_progress, self.selected_semester_id)
                )

                self.conn.commit()
                self.refresh_semester_goals()
                dialog.destroy()
                self.status_bar.config(text="Прогресс обновлён!")
                messagebox.showinfo("Успех", "Прогресс обновлён!")

            except ValueError:
                messagebox.showerror("Ошибка", "Введите число от 0 до 100")

        ttk.Button(dialog, text="Сохранить", command=save_progress, style="Primary.TButton").pack(pady=10)

    def delete_semester_goal(self):
        """Удаление цели на семестр"""
        if not hasattr(self, 'selected_semester_id') or not self.selected_semester_id:
            messagebox.showwarning("Внимание", "Выберите цель из списка")
            return

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            try:
                self.cursor.execute(
                    "DELETE FROM цели_на_семестр WHERE id = %s" if self.db_type == "postgres" else
                    "DELETE FROM цели_на_семестр WHERE id = ?",
                    (self.selected_semester_id,)
                )

                self.conn.commit()
                self.refresh_semester_goals()
                self.selected_semester_id = None
                self.status_bar.config(text="Цель удалена!")
                messagebox.showinfo("Успех", "Цель удалена!")

            except Exception as e:
                self.status_bar.config(text=f"Ошибка при удалении: {str(e)}")
                messagebox.showerror("Ошибка", f"Ошибка при удалении: {str(e)}")

    def refresh_semester_goals(self):
        """Обновление списка целей на семестр"""
        for item in self.semester_tree.get_children():
            self.semester_tree.delete(item)

        self.cursor.execute(
            "SELECT id, текст_цели, тип_цели, параметр, текущий_прогресс, целевой_прогресс FROM цели_на_семестр")
        goals = self.cursor.fetchall()

        for goal in goals:
            goal_id, text, goal_type, param, current, target = goal

            goal_text = text
            if param:
                goal_text += f" ({param})"

            if target > 0:
                percentage = (current / target) * 100
            else:
                percentage = 0

            progress_text = f"{current}/{target} ({percentage:.1f}%)"

            bar_length = 10
            filled = int((percentage / 100) * bar_length)
            progress_bar = "█" * filled + "░" * (bar_length - filled)

            self.semester_tree.insert("", tk.END, values=(goal_id, goal_text, progress_text, progress_bar))

    def setup_settings_tab(self):
        """Настройка вкладки Настройки"""
        main_frame = ttk.Frame(self.settings_frame, style="Custom.TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        ttk.Label(main_frame, text="Специальность:", font=('Arial', 11, 'bold')).pack(pady=20)

        self.specialty_var = tk.StringVar()
        specialties = ["Информационные системы", "Программная инженерия", "Прикладная информатика",
                       "Бизнес-информатика"]

        specialty_combo = ttk.Combobox(main_frame, textvariable=self.specialty_var, values=specialties,
                                       width=40, style="Custom.TCombobox", font=('Arial', 10))
        specialty_combo.pack(pady=10)

        ttk.Button(main_frame, text="💾 Сохранить", command=self.save_settings,
                   style="Primary.TButton").pack(pady=20)

        self.current_specialty_label = ttk.Label(main_frame,
                                                 text="Текущая специальность: не выбрана",
                                                 font=('Arial', 10),
                                                 foreground=self.colors['text_light'])
        self.current_specialty_label.pack(pady=10)

    def save_settings(self):
        """Сохранение настроек"""
        specialty = self.specialty_var.get()
        if specialty:
            self.current_specialty_label.config(text=f"Текущая специальность: {specialty}")
            self.status_bar.config(text=f"Специальность '{specialty}' сохранена")
            messagebox.showinfo("Успех", f"Специальность '{specialty}' сохранена!")
        else:
            messagebox.showwarning("Внимание", "Выберите специальность")

    def generate_report(self):
        """Формирование отчёта в формате Word"""
        try:
            self.status_bar.config(text="Формирование отчёта...")

            doc = Document()

            # Заголовок
            title = doc.add_heading('Индивидуальный образовательный маршрут', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(44, 111, 187)

            # Информация о студенте
            doc.add_heading('Общая информация', level=1)
            doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')

            # Цели
            doc.add_heading('Цели', level=1)

            self.cursor.execute("""
                SELECT название, тип, статус, план_дата, факт_дата, описание 
                FROM цели 
                ORDER BY план_дата
            """)
            goals = self.cursor.fetchall()

            for i, (name, goal_type, status, plan_date, fact_date, description) in enumerate(goals, 1):
                doc.add_heading(f'{i}. {name}', level=2)

                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Light Shading'

                rows = [
                    ("Тип цели:", goal_type),
                    ("Статус:", status),
                    ("Плановая дата:", plan_date or "не указана"),
                    ("Фактическая дата:", fact_date or "не указана")
                ]

                for j, (label, value) in enumerate(rows):
                    info_table.cell(j, 0).text = label
                    info_table.cell(j, 1).text = str(value)

                if description:
                    doc.add_paragraph('Описание:')
                    self.format_text_for_word(doc, description)

                doc.add_paragraph()

            # Навыки
            doc.add_heading('Навыки', level=1)

            self.cursor.execute("""
                SELECT н.название, COUNT(цн.цель_id) as количество_целей,
                       SUM(CASE WHEN ц.статус = 'Завершено' THEN 1 ELSE 0 END) as завершено_целей
                FROM навыки н
                LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
                LEFT JOIN цели ц ON цн.цель_id = ц.id
                GROUP BY н.id, н.название
                ORDER BY количество_целей DESC
            """)

            skills = self.cursor.fetchall()
            if skills:
                for skill, total, completed in skills:
                    p = doc.add_paragraph(f'• {skill} — всего целей: {total}', style='List Bullet')
                    if completed > 0:
                        p.add_run(f' (завершено: {completed})')
            else:
                doc.add_paragraph('Навыки не указаны')

            # Компетенции
            doc.add_heading('Компетенции', level=1)

            self.cursor.execute("""
                SELECT к.название, к.категория, 
                       ROUND(AVG(цк.уровень), 1) as средний_уровень,
                       COUNT(цк.уровень) as оценено_раз
                FROM компетенции к
                LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                GROUP BY к.id, к.название, к.категория
                ORDER BY средний_уровень DESC NULLS LAST
            """)

            comps = self.cursor.fetchall()

            if comps:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                table.autofit = False

                header_cells = table.rows[0].cells
                headers = ['Название', 'Категория', 'Средний уровень', 'Оценено раз']

                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].bold = True

                for name, category, avg_level, count in comps:
                    row_cells = table.add_row().cells
                    row_cells[0].text = name
                    row_cells[1].text = category or '—'
                    row_cells[2].text = f"{avg_level:.1f}" if avg_level else '—'
                    row_cells[3].text = str(count) if count else '—'

            # Слабые зоны
            doc.add_heading('Слабые зоны', level=1)
            weak_zones = [comp for comp in comps if comp[2] and comp[2] < 3]

            if weak_zones:
                for name, category, avg_level, _ in weak_zones:
                    doc.add_paragraph(f'⚠️ {name} ({category}): уровень {avg_level:.1f}', style='List Bullet')
            else:
                doc.add_paragraph('✅ Слабых зон не обнаружено')

            # Рекомендации
            doc.add_heading('Рекомендации', level=1)

            if weak_zones:
                for name, category, avg_level, _ in weak_zones[:3]:
                    if "презентация" in name.lower():
                        doc.add_paragraph(
                            f'• Для развития компетенции "{name}" рекомендуем выступить на студенческой конференции или подготовить доклад.')
                    elif "баз" in name.lower() or "данн" in name.lower():
                        doc.add_paragraph(
                            f'• Компетенция "{name}" требует внимания. Рассмотрите дополнительные курсы по базам данных или практические проекты.')
                    elif "проект" in name.lower() or "управление" in name.lower():
                        doc.add_paragraph(
                            f'• Для улучшения компетенции "{name}" участвуйте в командных проектах или возглавьте учебный проект.')
                    else:
                        doc.add_paragraph(
                            f'• Рекомендуем уделить больше внимания компетенции "{name}". Ищите практические задания по этой теме.')
            else:
                doc.add_paragraph('🎉 Все компетенции развиты хорошо. Продолжайте в том же духе!')

            # Достижения
            doc.add_heading('Достижения', level=1)

            self.cursor.execute("""
                SELECT название, описание 
                FROM достижения 
                WHERE получено = 1
                ORDER BY код
            """)
            achievements = self.cursor.fetchall()

            if achievements:
                for name, description in achievements:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run('🏆 ').bold = True
                    p.add_run(f'{name}: ').bold = True
                    p.add_run(description)
            else:
                doc.add_paragraph('Достижения пока не получены')

            # Цели на семестр
            doc.add_heading('Цели на семестр', level=1)

            self.cursor.execute("""
                SELECT текст_цели, тип_цели, параметр, текущий_прогресс, целевой_прогресс 
                FROM цели_на_семестр
                ORDER BY id
            """)
            semester_goals = self.cursor.fetchall()

            if semester_goals:
                for text, goal_type, param, current, target in semester_goals:
                    goal_text = f"{text}"
                    if param:
                        goal_text += f" ({param})"

                    if target > 0:
                        percentage = (current / target) * 100
                        progress_text = f" — {current} из {target} ({percentage:.1f}%)"

                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(goal_text)

                        progress_run = p.add_run(progress_text)
                        if percentage >= 100:
                            progress_run.font.color.rgb = RGBColor(16, 185, 129)
                        elif percentage >= 50:
                            progress_run.font.color.rgb = RGBColor(245, 158, 11)
                        else:
                            progress_run.font.color.rgb = RGBColor(239, 68, 68)
            else:
                doc.add_paragraph('Цели на семестр не установлены')

            # Сохранение файла
            filename = f'отчёт_ИОМ_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(filename)

            self.status_bar.config(text=f"Отчёт '{filename}' успешно создан!")
            messagebox.showinfo("Успех",
                                f"Отчёт успешно создан!\n\nФайл: {filename}\nРасположение: {os.path.abspath(filename)}",
                                icon='info')

        except Exception as e:
            self.status_bar.config(text="Ошибка создания отчёта")
            messagebox.showerror("Ошибка", f"Ошибка при создании отчёта: {str(e)}")

    def format_text_for_word(self, doc, text):
        """Форматирование текста для Word"""
        lines = text.split('\n')

        for line in lines:
            if not line.strip():
                doc.add_paragraph()
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('*') and line.endswith('*') and len(line) > 1:
                p = doc.add_paragraph()
                run = p.add_run(line[1:-1])
                run.italic = True
            elif '[' in line and '](' in line and ')' in line:
                start = line.find('[') + 1
                end = line.find(']')
                link_text = line[start:end]

                url_start = line.find('](') + 2
                url_end = line.find(')', url_start)
                url = line[url_start:url_end]

                paragraph = doc.add_paragraph()
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                paragraph.add_run(f" ({url})")
            else:
                doc.add_paragraph(line)

    def __del__(self):
        """Закрытие соединения с БД при завершении"""
        if hasattr(self, 'conn'):
            try:
                self.conn.close()
            except:
                pass


def main():
    """Основная функция запуска приложения"""
    # Создание файла компетенций по умолчанию если его нет
    if not os.path.exists('competencies.json'):
        default_competencies = [
            {"название": "Работа с БД", "категория": "Технические"},
            {"название": "Презентация результатов", "категория": "Коммуникативные"},
            {"название": "Управление проектами", "категория": "Организационные"},
            {"название": "Анализ данных", "категория": "Технические"},
            {"название": "Программирование", "категория": "Технические"},
            {"название": "Командная работа", "категория": "Коммуникативные"},
            {"название": "Самоорганизация", "категория": "Организационные"}
        ]
        with open('competencies.json', 'w', encoding='utf-8') as f:
            json.dump(default_competencies, f, ensure_ascii=False, indent=2)
        print("Создан файл competencies.json с компетенциями по умолчанию")

    # Запуск приложения
    root = tk.Tk()
    app = EducationalRoutePlanner(root)

    # Обработка закрытия окна
    def on_closing():
        if messagebox.askokcancel("Выход", "Вы уверены, что хотите выйти?"):
            root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()


if __name__ == "__main__":
    main()