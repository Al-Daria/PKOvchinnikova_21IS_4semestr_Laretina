import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
from datetime import datetime
from pathlib import Path
import os
import sys
import threading
import webbrowser
import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl.styles import Font as ExcelFont, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Импортируем менеджер БД
from database_manager import DatabaseManager


class ResearchPortfolioApp:
    """Главное приложение электронного портфолио"""

    def __init__(self, root):
        self.root = root
        self.root.title("Электронный портфолио исследователя")
        self.root.geometry("1400x800")
        self.center_window()

        # Инициализация БД
        self.db = DatabaseManager()
        if not self.db.connection:
            messagebox.showerror("Ошибка", "Не удалось подключиться к базе данных!")
            sys.exit(1)

        # Текущая запись
        self.current_entry_id = None
        self.current_file_path = None

        # Цвета
        self.colors = {
            'primary': '#2C3E50',
            'secondary': '#3498DB',
            'success': '#27AE60',
            'danger': '#E74C3C',
            'light': '#ECF0F1'
        }

        # Создание интерфейса
        self.create_interface()

        # Загрузка записей
        self.load_entries()

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def center_window(self):
        """Центрирование окна"""
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')

    def create_interface(self):
        """Создание интерфейса"""
        # Notebook (вкладки)
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Вкладки
        self.create_tab()
        self.edit_tab()
        self.analytics_tab()

        # Статус бар
        self.status_bar = tk.Label(self.root, text="Готово", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def create_tab(self):
        """Вкладка создания записей"""
        tab = ttk.Frame(self.notebook)

        # Заголовок
        tk.Label(tab, text="Создание новой записи", font=('Arial', 14, 'bold'),
                 fg=self.colors['primary']).pack(pady=20)

        # Основной фрейм
        main_frame = ttk.Frame(tab)
        main_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Левая панель
        left_frame = ttk.LabelFrame(main_frame, text="Основная информация", padding=15)
        left_frame.pack(side=tk.LEFT, fill='both', expand=True, padx=(0, 10))

        # Поля формы
        row = 0

        tk.Label(left_frame, text="Название*:").grid(row=row, column=0, sticky='w', pady=(0, 10))
        self.title_entry = ttk.Entry(left_frame, width=40)
        self.title_entry.grid(row=row, column=1, pady=(0, 10), padx=(10, 0))
        row += 1

        tk.Label(left_frame, text="Тип*:").grid(row=row, column=0, sticky='w', pady=10)
        self.type_combo = ttk.Combobox(left_frame, values=self.db.ENTRY_TYPES,
                                       state='readonly', width=37)
        self.type_combo.grid(row=row, column=1, pady=10, padx=(10, 0))
        self.type_combo.current(0)
        row += 1

        tk.Label(left_frame, text="Год*:").grid(row=row, column=0, sticky='w', pady=10)
        self.year_entry = ttk.Entry(left_frame, width=40)
        self.year_entry.grid(row=row, column=1, pady=10, padx=(10, 0))
        row += 1

        # Соавторы
        tk.Label(left_frame, text="Соавторы:").grid(row=row, column=0, sticky='w', pady=(20, 5))
        row += 1

        coauthor_frame = ttk.Frame(left_frame)
        coauthor_frame.grid(row=row, column=0, columnspan=2, sticky='ew', pady=(0, 10))

        self.coauthor_entry = ttk.Entry(coauthor_frame, width=30)
        self.coauthor_entry.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(coauthor_frame, text="Добавить",
                   command=self.add_coauthor).pack(side=tk.LEFT)
        row += 1

        # Список соавторов
        self.coauthors_listbox = tk.Listbox(left_frame, height=6)
        self.coauthors_listbox.grid(row=row, column=0, columnspan=2, sticky='ew', pady=(0, 10))
        row += 1

        # Кнопки соавторов
        coauthor_buttons = ttk.Frame(left_frame)
        coauthor_buttons.grid(row=row, column=0, columnspan=2, pady=(0, 20))

        ttk.Button(coauthor_buttons, text="Удалить выбранного",
                   command=self.remove_coauthor).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(coauthor_buttons, text="Очистить список",
                   command=self.clear_coauthors).pack(side=tk.LEFT)

        # Правая панель - описание
        right_frame = ttk.LabelFrame(main_frame, text="Описание (Markdown)", padding=15)
        right_frame.pack(side=tk.RIGHT, fill='both', expand=True, padx=(10, 0))

        # Текстовое поле с скроллбаром
        text_frame = ttk.Frame(right_frame)
        text_frame.pack(fill='both', expand=True)

        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.description_text = tk.Text(text_frame, height=20, width=60,
                                        font=('Courier', 10), wrap=tk.WORD,
                                        yscrollcommand=scrollbar.set)
        self.description_text.pack(side=tk.LEFT, fill='both', expand=True)
        scrollbar.config(command=self.description_text.yview)

        # Подсказка по Markdown
        help_text = "• **Жирный текст**\n• *Курсив*\n• [Ссылка](url)\n• > Цитата\n• ```код```\n• # Заголовок"
        tk.Label(right_frame, text=help_text, justify=tk.LEFT,
                 anchor='w').pack(fill='x', pady=(10, 0))

        # Кнопка создания
        create_btn = tk.Button(tab, text="СОЗДАТЬ ЗАПИСЬ",
                               font=('Arial', 12, 'bold'),
                               bg=self.colors['success'], fg='white',
                               command=self.create_entry,
                               padx=20, pady=10, cursor='hand2')
        create_btn.pack(pady=20)

        self.notebook.add(tab, text='➕ Создание')

    def edit_tab(self):
        """Вкладка редактирования"""
        tab = ttk.Frame(self.notebook)

        # Панель с разделителем
        paned = ttk.PanedWindow(tab, orient=tk.HORIZONTAL)
        paned.pack(fill='both', expand=True, padx=10, pady=10)

        # Левая панель - список записей
        left_panel = ttk.LabelFrame(paned, text="Список записей", padding=10)

        # Панель инструментов
        toolbar = ttk.Frame(left_panel)
        toolbar.pack(fill='x', pady=(0, 10))

        ttk.Button(toolbar, text="Обновить",
                   command=self.load_entries).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(toolbar, text="Сортировка",
                   command=self.sort_entries).pack(side=tk.LEFT)

        # Treeview
        tree_frame = ttk.Frame(left_panel)
        tree_frame.pack(fill='both', expand=True)

        scrollbar = ttk.Scrollbar(tree_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.tree = ttk.Treeview(tree_frame, yscrollcommand=scrollbar.set,
                                 selectmode='browse', height=20)
        scrollbar.config(command=self.tree.yview)

        # Колонки
        self.tree['columns'] = ('ID', 'Название', 'Тип', 'Год', 'Создано')
        self.tree.column('#0', width=0, stretch=tk.NO)
        self.tree.column('ID', width=50, anchor=tk.CENTER)
        self.tree.column('Название', width=300, anchor=tk.W)
        self.tree.column('Тип', width=120, anchor=tk.W)
        self.tree.column('Год', width=80, anchor=tk.CENTER)
        self.tree.column('Создано', width=150, anchor=tk.CENTER)

        # Заголовки
        for col in self.tree['columns']:
            self.tree.heading(col, text=col, anchor=tk.CENTER if col == 'ID' or col == 'Год' else tk.W)

        self.tree.pack(fill='both', expand=True)
        self.tree.bind('<<TreeviewSelect>>', self.on_tree_select)

        paned.add(left_panel, weight=1)

        # Правая панель - редактирование
        right_panel = ttk.LabelFrame(paned, text="Редактирование записи", padding=15)

        edit_frame = ttk.Frame(right_panel)
        edit_frame.pack(fill='both', expand=True)

        row = 0

        tk.Label(edit_frame, text="Название:").grid(row=row, column=0, sticky='w', pady=(0, 10))
        self.edit_title = ttk.Entry(edit_frame, width=40)
        self.edit_title.grid(row=row, column=1, pady=(0, 10), padx=(10, 0))
        row += 1

        tk.Label(edit_frame, text="Тип:").grid(row=row, column=0, sticky='w', pady=10)
        self.edit_type = ttk.Combobox(edit_frame, values=self.db.ENTRY_TYPES,
                                      state='readonly', width=37)
        self.edit_type.grid(row=row, column=1, pady=10, padx=(10, 0))
        row += 1

        tk.Label(edit_frame, text="Год:").grid(row=row, column=0, sticky='w', pady=10)
        self.edit_year = ttk.Entry(edit_frame, width=40)
        self.edit_year.grid(row=row, column=1, pady=10, padx=(10, 0))
        row += 1

        tk.Label(edit_frame, text="Описание:").grid(row=row, column=0, sticky='nw', pady=(10, 0))

        # Текстовое поле с скроллбаром
        text_frame = ttk.Frame(edit_frame)
        text_frame.grid(row=row, column=1, sticky='nsew', pady=(10, 0), padx=(10, 0))

        text_scroll = ttk.Scrollbar(text_frame)
        text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.edit_description = tk.Text(text_frame, height=15, width=50,
                                        font=('Courier', 10), wrap=tk.WORD,
                                        yscrollcommand=text_scroll.set)
        self.edit_description.pack(side=tk.LEFT, fill='both', expand=True)
        text_scroll.config(command=self.edit_description.yview)

        edit_frame.grid_rowconfigure(row, weight=1)
        edit_frame.grid_columnconfigure(1, weight=1)

        # Кнопки управления
        button_frame = ttk.Frame(right_panel)
        button_frame.pack(fill='x', pady=(20, 0))

        self.save_btn = ttk.Button(button_frame, text="Сохранить",
                                   command=self.save_entry, state='disabled')
        self.save_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.delete_btn = ttk.Button(button_frame, text="Удалить",
                                     command=self.delete_entry, state='disabled')
        self.delete_btn.pack(side=tk.LEFT, padx=(0, 10))

        self.open_file_btn = ttk.Button(button_frame, text="Открыть файл",
                                        command=self.open_file, state='disabled')
        self.open_file_btn.pack(side=tk.LEFT)

        paned.add(right_panel, weight=1)

        self.notebook.add(tab, text='✏️ Редактирование')

    def analytics_tab(self):
        """Вкладка аналитики"""
        tab = ttk.Frame(self.notebook)

        tk.Label(tab, text="Аналитика и отчётность",
                 font=('Arial', 14, 'bold'), fg=self.colors['primary']).pack(pady=30)

        # Кнопка генерации отчета
        report_btn = tk.Button(tab, text="📊 СФОРМИРОВАТЬ ОТЧЕТ",
                               font=('Arial', 14, 'bold'),
                               bg=self.colors['secondary'], fg='white',
                               command=self.generate_report,
                               padx=30, pady=15, cursor='hand2')
        report_btn.pack(pady=20)

        # Область статистики
        stats_frame = ttk.LabelFrame(tab, text="Статистика портфолио", padding=20)
        stats_frame.pack(fill='both', expand=True, padx=20, pady=20)

        self.stats_text = tk.Text(stats_frame, height=20, width=80,
                                  font=('Arial', 10), wrap=tk.WORD)
        self.stats_text.pack(fill='both', expand=True)
        self.stats_text.insert('1.0',
                               "Для просмотра статистики нажмите кнопку 'Сформировать отчет'.\n\n"
                               "После генерации отчета здесь будет отображена:\n"
                               "• Общая статистика по записям\n"
                               "• Распределение по типам\n"
                               "• Динамика по годам\n"
                               "• Количество уникальных соавторов\n"
                               "• Список последних записей")
        self.stats_text.config(state='disabled')

        self.notebook.add(tab, text='📊 Аналитика')

    def load_entries(self):
        """Загрузка записей"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        entries = self.db.get_entries()

        for entry in entries:
            self.tree.insert('', 'end', values=entry)

        self.update_status(f"Загружено записей: {len(entries)}")

    def on_tree_select(self, event):
        """Обработка выбора записи"""
        selected = self.tree.selection()
        if not selected:
            return

        item = selected[0]
        values = self.tree.item(item, 'values')

        if values:
            self.current_entry_id = int(values[0])

            self.edit_title.delete(0, tk.END)
            self.edit_title.insert(0, values[1])

            self.edit_type.set(values[2])

            self.edit_year.delete(0, tk.END)
            self.edit_year.insert(0, values[3])

            # Получаем путь к файлу
            cursor = self.db.connection.cursor()
            cursor.execute("SELECT file_path FROM entries WHERE id = %s", (self.current_entry_id,))
            result = cursor.fetchone()
            cursor.close()

            if result:
                self.current_file_path = result[0]

                if os.path.exists(self.current_file_path):
                    with open(self.current_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        self.edit_description.delete('1.0', tk.END)
                        self.edit_description.insert('1.0', content)

            self.save_btn.config(state='normal')
            self.delete_btn.config(state='normal')
            self.open_file_btn.config(state='normal')

    def create_entry(self):
        """Создание записи"""
        title = self.title_entry.get().strip()
        entry_type = self.type_combo.get()
        year = self.year_entry.get().strip()
        description = self.description_text.get('1.0', tk.END).strip()

        if not title:
            messagebox.showwarning("Предупреждение", "Введите название!")
            self.title_entry.focus()
            return

        if not year:
            messagebox.showwarning("Предупреждение", "Введите год!")
            self.year_entry.focus()
            return

        try:
            year_int = int(year)
            current_year = datetime.now().year
            if year_int < 1900 or year_int > current_year + 1:
                raise ValueError
        except ValueError:
            messagebox.showwarning("Ошибка", "Введите корректный год!")
            return

        # Создаем папку и файл
        files_dir = Path("portfolio_files")
        files_dir.mkdir(exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = ''.join(c if c.isalnum() else '_' for c in title)[:30]
        filename = f"{timestamp}_{safe_title}.md"
        file_path = files_dir / filename

        try:
            entry_id = self.db.create_entry(title, entry_type, year_int, str(file_path))

            # Добавляем соавторов
            for i in range(self.coauthors_listbox.size()):
                coauthor = self.coauthors_listbox.get(i)
                if coauthor.strip():
                    self.db.add_coauthor(entry_id, coauthor)

            # Сохраняем файл
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"# {title}\n\n")
                f.write(f"**Тип:** {entry_type}\n")
                f.write(f"**Год:** {year}\n")
                f.write(f"**Дата:** {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n")

                if self.coauthors_listbox.size() > 0:
                    f.write("**Соавторы:**\n")
                    for i in range(self.coauthors_listbox.size()):
                        f.write(f"- {self.coauthors_listbox.get(i)}\n")
                    f.write("\n")

                f.write("## Описание\n\n")
                f.write(description if description else "Описание отсутствует")

            # Очищаем форму
            self.title_entry.delete(0, tk.END)
            self.year_entry.delete(0, tk.END)
            self.description_text.delete('1.0', tk.END)
            self.coauthors_listbox.delete(0, tk.END)

            self.load_entries()

            messagebox.showinfo("Успех", f"Запись создана! ID: {entry_id}")
            self.update_status(f"Создана: {title}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка:\n{str(e)}")

    def add_coauthor(self):
        """Добавление соавтора"""
        coauthor = self.coauthor_entry.get().strip()
        if coauthor:
            existing = self.coauthors_listbox.get(0, tk.END)
            if coauthor not in existing:
                self.coauthors_listbox.insert(tk.END, coauthor)
                self.coauthor_entry.delete(0, tk.END)
            else:
                messagebox.showwarning("Предупреждение", "Соавтор уже добавлен!")
        else:
            messagebox.showwarning("Предупреждение", "Введите ФИО соавтора!")

    def remove_coauthor(self):
        """Удаление соавтора"""
        selected = self.coauthors_listbox.curselection()
        if selected:
            self.coauthors_listbox.delete(selected[0])
        else:
            messagebox.showwarning("Предупреждение", "Выберите соавтора!")

    def clear_coauthors(self):
        """Очистка списка соавторов"""
        if self.coauthors_listbox.size() > 0:
            if messagebox.askyesno("Подтверждение", "Очистить список соавторов?"):
                self.coauthors_listbox.delete(0, tk.END)

    def save_entry(self):
        """Сохранение записи"""
        if not self.current_entry_id:
            return

        title = self.edit_title.get().strip()
        entry_type = self.edit_type.get()
        year = self.edit_year.get().strip()
        description = self.edit_description.get('1.0', tk.END).strip()

        if not title or not year:
            messagebox.showwarning("Предупреждение", "Заполните все поля!")
            return

        try:
            year_int = int(year)
        except ValueError:
            messagebox.showwarning("Ошибка", "Год должен быть числом!")
            return

        try:
            self.db.update_entry(self.current_entry_id, title, entry_type, year_int)

            if self.current_file_path and os.path.exists(self.current_file_path):
                with open(self.current_file_path, 'w', encoding='utf-8') as f:
                    f.write(description)

            self.load_entries()
            messagebox.showinfo("Успех", "Изменения сохранены!")
            self.update_status(f"Обновлена: {title}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка:\n{str(e)}")

    def delete_entry(self):
        """Удаление записи"""
        if not self.current_entry_id:
            return

        if not messagebox.askyesno("Подтверждение", "Удалить запись безвозвратно?"):
            return

        try:
            if self.current_file_path and os.path.exists(self.current_file_path):
                os.remove(self.current_file_path)

            self.db.delete_entry(self.current_entry_id)

            self.edit_title.delete(0, tk.END)
            self.edit_year.delete(0, tk.END)
            self.edit_description.delete('1.0', tk.END)

            self.save_btn.config(state='disabled')
            self.delete_btn.config(state='disabled')
            self.open_file_btn.config(state='disabled')

            self.load_entries()
            messagebox.showinfo("Успех", "Запись удалена!")
            self.update_status("Запись удалена")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка:\n{str(e)}")

    def open_file(self):
        """Открытие файла"""
        if self.current_file_path and os.path.exists(self.current_file_path):
            try:
                webbrowser.open(f'file://{os.path.abspath(self.current_file_path)}')
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{str(e)}")
        else:
            messagebox.showwarning("Ошибка", "Файл не найден!")

    def sort_entries(self):
        """Сортировка записей"""
        sort_options = [
            "Дата создания (новые)",
            "Дата создания (старые)",
            "Название (А-Я)",
            "Название (Я-А)",
            "Год (по возрастанию)",
            "Год (по убыванию)"
        ]

        sort_dialog = tk.Toplevel(self.root)
        sort_dialog.title("Сортировка")
        sort_dialog.geometry("300x200")
        sort_dialog.transient(self.root)
        sort_dialog.grab_set()

        sort_dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (300 // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (200 // 2)
        sort_dialog.geometry(f"+{x}+{y}")

        tk.Label(sort_dialog, text="Выберите тип сортировки:").pack(pady=20)

        sort_var = tk.StringVar(value=sort_options[0])

        for option in sort_options:
            ttk.Radiobutton(sort_dialog, text=option, variable=sort_var,
                            value=option).pack(anchor='w', padx=20)

        def apply_sort():
            selected = sort_var.get()

            sort_map = {
                "Дата создания (новые)": ("created_at", "DESC"),
                "Дата создания (старые)": ("created_at", "ASC"),
                "Название (А-Я)": ("title", "ASC"),
                "Название (Я-А)": ("title", "DESC"),
                "Год (по возрастанию)": ("year", "ASC"),
                "Год (по убыванию)": ("year", "DESC")
            }

            if selected in sort_map:
                sort_by, sort_order = sort_map[selected]

                for item in self.tree.get_children():
                    self.tree.delete(item)

                entries = self.db.get_entries(sort_by, sort_order)
                for entry in entries:
                    self.tree.insert('', 'end', values=entry)

                self.update_status(f"Сортировка: {selected}")

            sort_dialog.destroy()

        button_frame = ttk.Frame(sort_dialog)
        button_frame.pack(pady=20)

        ttk.Button(button_frame, text="Применить",
                   command=apply_sort).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(button_frame, text="Отмена",
                   command=sort_dialog.destroy).pack(side=tk.LEFT)

    def generate_report(self):
        """Генерация отчетов"""
        entries = self.db.get_entries()
        if not entries:
            if not messagebox.askyesno("Нет данных",
                                       "В базе нет записей. Создать отчет с нулевыми данными?"):
                return

        # Создаем диалог прогресса
        progress = tk.Toplevel(self.root)
        progress.title("Генерация отчетов")
        progress.geometry("400x150")
        progress.transient(self.root)
        progress.grab_set()

        progress.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (400 // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (150 // 2)
        progress.geometry(f"+{x}+{y}")

        tk.Label(progress, text="Идет генерация отчетов...").pack(pady=20)

        progress_var = tk.IntVar()
        progress_bar = ttk.Progressbar(progress, variable=progress_var, maximum=100)
        progress_bar.pack(fill='x', padx=20, pady=10)

        status_label = tk.Label(progress, text="Подготовка...")
        status_label.pack()

        def update_progress(value, status):
            progress_var.set(value)
            status_label.config(text=status)
            progress.update()

        def generate_in_thread():
            try:
                update_progress(10, "Получение статистики...")

                # Получаем статистику
                stats = self.db.get_statistics()

                update_progress(30, "Создание папок...")

                # Создаем папки
                reports_dir = Path("reports")
                reports_dir.mkdir(exist_ok=True)

                update_progress(40, "Создание графиков...")

                # Создаем графики
                self.create_charts(stats)

                update_progress(60, "Генерация Excel...")

                # Генерируем Excel
                excel_path = self.generate_excel(stats)

                update_progress(80, "Генерация Word...")

                # Генерируем Word
                word_path = self.generate_word(stats)

                update_progress(100, "Завершение...")

                # Обновляем статистику в интерфейсе
                self.display_statistics(stats)

                progress.destroy()

                messagebox.showinfo("Успех",
                                    f"✅ Отчеты созданы!\n\n"
                                    f"Excel: {excel_path}\n"
                                    f"Word: {word_path}")

                self.update_status("Отчеты сгенерированы")

            except Exception as e:
                progress.destroy()
                messagebox.showerror("Ошибка", f"Ошибка:\n{str(e)}")

        thread = threading.Thread(target=generate_in_thread)
        thread.daemon = True
        thread.start()

        self.root.after(100, lambda: self.check_thread(thread, progress))

    def check_thread(self, thread, progress):
        """Проверка потока"""
        if thread.is_alive():
            self.root.after(100, lambda: self.check_thread(thread, progress))

    def create_charts(self, stats):
        """Создание графиков"""
        reports_dir = Path("reports")

        # График распределения по типам
        if stats['type_distribution']:
            plt.figure(figsize=(10, 6))
            types = list(stats['type_distribution'].keys())
            counts = list(stats['type_distribution'].values())

            colors = ['#4CAF50', '#2196F3', '#FF9800', '#9C27B0', '#F44336']
            bars = plt.bar(types, counts, color=colors[:len(types)], edgecolor='black')

            plt.title('Распределение записей по типам', fontsize=14, fontweight='bold')
            plt.xlabel('Тип записи')
            plt.ylabel('Количество')
            plt.xticks(rotation=45, ha='right')
            plt.grid(axis='y', alpha=0.3)

            for bar, count in zip(bars, counts):
                plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                         str(count), ha='center', va='bottom')

            plt.tight_layout()
            plt.savefig(reports_dir / 'type_distribution.png', dpi=300)
            plt.close()

        # График распределения по годам
        if stats['year_distribution']:
            plt.figure(figsize=(12, 6))
            years = sorted(stats['year_distribution'].keys())
            counts = [stats['year_distribution'][y] for y in years]

            plt.bar([str(y) for y in years], counts, color='#2196F3', edgecolor='black')
            plt.title('Динамика по годам', fontsize=14, fontweight='bold')
            plt.xlabel('Год')
            plt.ylabel('Количество записей')
            plt.grid(axis='y', alpha=0.3)

            plt.tight_layout()
            plt.savefig(reports_dir / 'year_distribution.png', dpi=300)
            plt.close()

    def generate_excel(self, stats):
        """Генерация Excel отчета"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Статистика"

        # Заголовок
        ws.merge_cells('A1:D1')
        title_cell = ws['A1']
        title_cell.value = "Отчет по портфолио исследователя"
        title_cell.font = ExcelFont(bold=True, size=16)
        title_cell.alignment = Alignment(horizontal='center')

        ws['A2'] = f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}"

        # Ключевые показатели
        ws['A4'] = "Ключевые показатели:"
        ws['A4'].font = ExcelFont(bold=True)

        data = [
            ["Показатель", "Значение"],
            ["Всего записей", stats['total_entries']],
            ["Уникальных соавторов", stats['unique_coauthors']],
            ["Типов записей", len(stats['type_distribution'])],
        ]

        for i, row in enumerate(data, start=5):
            for j, value in enumerate(row, start=1):
                cell = ws.cell(row=i, column=j, value=value)
                if i == 5:
                    cell.font = ExcelFont(bold=True)
                    cell.fill = PatternFill(start_color="DDDDDD", fill_type="solid")

        # Распределение по типам
        start_row = 10
        ws.cell(row=start_row, column=1, value="Распределение по типам:").font = ExcelFont(bold=True)

        if stats['type_distribution']:
            row = start_row + 1
            for entry_type, count in stats['type_distribution'].items():
                ws.cell(row=row, column=1, value=entry_type)
                ws.cell(row=row, column=2, value=count)
                row += 1

        # Лист с графиками
        ws2 = wb.create_sheet("Графики")

        if os.path.exists("reports/type_distribution.png"):
            img = ExcelImage("reports/type_distribution.png")
            img.width = 500
            img.height = 300
            ws2.add_image(img, 'A1')

        if os.path.exists("reports/year_distribution.png"):
            img = ExcelImage("reports/year_distribution.png")
            img.width = 500
            img.height = 300
            ws2.add_image(img, 'A20')

        # Сохраняем
        excel_path = "reports/portfolio_report.xlsx"
        wb.save(excel_path)

        return excel_path

    def generate_word(self, stats):
        """Генерация Word отчета"""
        doc = Document()

        # Стили
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Титульный лист
        title = doc.add_paragraph()
        title_run = title.add_run('ОТЧЕТ\nпо портфолио исследователя')
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('\n')

        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(14)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Ключевые показатели
        doc.add_heading('Ключевые показатели', level=1)

        table_data = [
            ['Показатель', 'Значение'],
            ['Всего записей', str(stats['total_entries'])],
            ['Уникальных соавторов', str(stats['unique_coauthors'])],
            ['Типов записей', str(len(stats['type_distribution']))],
        ]

        table = doc.add_table(rows=4, cols=2)
        table.style = 'LightShading'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)

        # Графики
        doc.add_heading('Визуализация данных', level=1)

        if os.path.exists("reports/type_distribution.png"):
            doc.add_paragraph('Распределение по типам:')
            doc.add_picture("reports/type_distribution.png", width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        if os.path.exists("reports/year_distribution.png"):
            doc.add_paragraph('Динамика по годам:')
            doc.add_picture("reports/year_distribution.png", width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Последние записи
        doc.add_page_break()
        doc.add_heading('Последние записи', level=1)

        if stats['recent_entries']:
            table = doc.add_table(rows=len(stats['recent_entries']) + 1, cols=4)
            table.style = 'LightShading'

            headers = ['Название', 'Тип', 'Год', 'Дата создания']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].font.bold = True

            for i, entry in enumerate(stats['recent_entries'], start=1):
                for j, value in enumerate(entry):
                    table.cell(i, j).text = str(value) if value is not None else ""

        # Сохраняем
        word_path = "reports/portfolio_report.docx"
        doc.save(word_path)

        return word_path

    def display_statistics(self, stats):
        """Отображение статистики"""
        self.stats_text.config(state='normal')
        self.stats_text.delete('1.0', tk.END)

        text = "📊 СТАТИСТИКА ПОРТФОЛИО\n\n"
        text += f"Всего записей: {stats['total_entries']}\n"
        text += f"Уникальных соавторов: {stats['unique_coauthors']}\n\n"

        if stats['type_distribution']:
            text += "Распределение по типам:\n"
            for entry_type, count in stats['type_distribution'].items():
                text += f"• {entry_type}: {count}\n"
            text += "\n"

        if stats['year_distribution']:
            text += "Распределение по годам:\n"
            for year, count in sorted(stats['year_distribution'].items()):
                text += f"• {year}: {count}\n"
            text += "\n"

        if stats['recent_entries']:
            text += "Последние записи:\n"
            for entry in stats['recent_entries']:
                text += f"• {entry[0]} ({entry[1]}, {entry[2]})\n"

        self.stats_text.insert('1.0', text)
        self.stats_text.config(state='disabled')

    def update_status(self, message):
        """Обновление статуса"""
        self.status_bar.config(text=message)

    def on_closing(self):
        """Обработка закрытия"""
        if messagebox.askokcancel("Выход", "Вы уверены, что хотите выйти?"):
            if hasattr(self, 'db'):
                self.db.close()
            self.root.destroy()


def main():
    """Запуск приложения"""
    root = tk.Tk()
    app = ResearchPortfolioApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()