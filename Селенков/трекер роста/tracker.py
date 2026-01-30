import tkinter as tk
from tkinter import ttk, messagebox
import psycopg2
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

DB_LOGIN = {
    'host': 'localhost',
    'database': 'postgres',
    'user': 'postgres',
    'password': '1111',
    'port': '5432'
}


class PortfolioApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Система портфолио и компетенций")
        self.root.geometry("1200x700")

        self.current_user_id = 1
        self.existing_keywords = []

        try:
            self.conn = psycopg2.connect(**DB_LOGIN)
            self.cursor = self.conn.cursor()
            self.initialize_database()
            self.load_competencies_from_json()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось подключиться к БД: {str(e)}")
            self.root.destroy()
            return

        self.create_widgets()
        self.load_entries()
        self.update_statistics()

    def initialize_database(self):
        tables = [
            '''CREATE TABLE IF NOT EXISTS entries (
                id SERIAL PRIMARY KEY,
                title TEXT NOT NULL,
                type TEXT NOT NULL,
                date DATE NOT NULL,
                description TEXT,
                coauthors TEXT,
                user_id INTEGER
            )''',
            '''CREATE TABLE IF NOT EXISTS keywords (
                id SERIAL PRIMARY KEY,
                keyword TEXT UNIQUE NOT NULL
            )''',
            '''CREATE TABLE IF NOT EXISTS entry_keywords (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                keyword_id INTEGER REFERENCES keywords(id) ON DELETE CASCADE,
                PRIMARY KEY (entry_id, keyword_id)
            )''',
            '''CREATE TABLE IF NOT EXISTS achievements (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                description TEXT,
                user_id INTEGER,
                unlocked_date DATE
            )''',
            '''CREATE TABLE IF NOT EXISTS competencies (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                category TEXT
            )''',
            '''CREATE TABLE IF NOT EXISTS entry_competencies (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                competency_id INTEGER REFERENCES competencies(id) ON DELETE CASCADE,
                level INTEGER CHECK (level >= 1 AND level <= 5),
                PRIMARY KEY (entry_id, competency_id)
            )''',
            '''CREATE TABLE IF NOT EXISTS goals (
                id SERIAL PRIMARY KEY,
                description TEXT NOT NULL,
                target_value INTEGER,
                current_value INTEGER,
                deadline DATE,
                user_id INTEGER
            )'''
        ]

        for table in tables:
            self.cursor.execute(table)
        self.conn.commit()

    def load_competencies_from_json(self):
        try:
            if os.path.exists('competencies.json'):
                with open('competencies.json', 'r', encoding='utf-8') as f:
                    data = json.load(f)

                self.cursor.execute("DELETE FROM competencies")

                for specialty_data in data:
                    for comp in specialty_data.get('competencies', []):
                        self.cursor.execute(
                            "INSERT INTO competencies (name, category) VALUES (%s, %s)",
                            (comp['name'], comp['category'])
                        )

                self.conn.commit()
            else:
                self.create_default_json()
                self.load_competencies_from_json()

        except Exception as e:
            print(f"Ошибка загрузки JSON: {e}")
            self.load_default_competencies()

    def create_default_json(self):
        default_data = [{
            "specialty": "Информационные системы",
            "competencies": [
                {"name": "Программирование", "category": "Технические"},
                {"name": "Работа с БД", "category": "Технические"},
                {"name": "Анализ данных", "category": "Технические"},
                {"name": "Проектная деятельность", "category": "Профессиональные"},
                {"name": "Научная работа", "category": "Профессиональные"},
                {"name": "Презентация результатов", "category": "Коммуникативные"},
                {"name": "Командная работа", "category": "Коммуникативные"},
                {"name": "Самоорганизация", "category": "Личные"}
            ]
        }]

        with open('competencies.json', 'w', encoding='utf-8') as f:
            json.dump(default_data, f, ensure_ascii=False, indent=2)

    def load_default_competencies(self):
        default = [
            ("Программирование", "Технические"),
            ("Работа с БД", "Технические"),
            ("Анализ данных", "Технические"),
            ("Проектная деятельность", "Профессиональные"),
            ("Научная работа", "Профессиональные"),
            ("Презентация результатов", "Коммуникативные"),
            ("Командная работа", "Коммуникативные"),
            ("Самоорганизация", "Личные")
        ]

        for name, category in default:
            self.cursor.execute(
                "INSERT INTO competencies (name, category) VALUES (%s, %s) ON CONFLICT DO NOTHING",
                (name, category)
            )
        self.conn.commit()

    def create_widgets(self):
        # Меню
        menubar = tk.Menu(self.root)
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Экспорт в Word", command=self.export_to_word)
        file_menu.add_separator()
        file_menu.add_command(label="Обновить", command=self.update_statistics)
        file_menu.add_command(label="Выход", command=self.root.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)
        self.root.config(menu=menubar)

        # Вкладки
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.create_add_entry_tab()
        self.create_view_tab()
        self.create_research_map_tab()
        self.create_achievements_tab()
        self.create_competencies_tab()
        self.create_goals_tab()

    def create_add_entry_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Добавить запись")

        # Название
        tk.Label(frame, text="Название:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        self.title_entry = tk.Entry(frame, width=50)
        self.title_entry.grid(row=0, column=1, padx=10, pady=5)

        # Тип
        tk.Label(frame, text="Тип:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        self.type_combo = ttk.Combobox(frame, values=["Проект", "Публикация", "Конференция", "Практика", "Грант"],
                                       state="readonly")
        self.type_combo.grid(row=1, column=1, padx=10, pady=5)

        # Дата
        tk.Label(frame, text="Дата (ГГГГ-ММ-ДД):").grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        self.date_entry = tk.Entry(frame, width=20)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.date_entry.grid(row=2, column=1, padx=10, pady=5)

        # Описание
        tk.Label(frame, text="Описание:").grid(row=3, column=0, sticky=tk.NW, padx=10, pady=5)
        self.description_text = tk.Text(frame, width=50, height=5)
        self.description_text.grid(row=3, column=1, padx=10, pady=5)

        # Соавторы
        tk.Label(frame, text="Соавторы (через запятую):").grid(row=4, column=0, sticky=tk.W, padx=10, pady=5)
        self.coauthors_entry = tk.Entry(frame, width=50)
        self.coauthors_entry.grid(row=4, column=1, padx=10, pady=5)

        # Ключевые слова
        tk.Label(frame, text="Ключевые слова (до 5, автодополнение):").grid(row=5, column=0, sticky=tk.W, padx=10,
                                                                            pady=5)
        self.keywords_frame = tk.Frame(frame)
        self.keywords_frame.grid(row=5, column=1, padx=10, pady=5)

        self.keyword_combos = []
        for i in range(5):
            combo = ttk.Combobox(self.keywords_frame, width=15)
            combo.grid(row=0, column=i, padx=2)
            combo.bind('<KeyRelease>', lambda e, idx=i: self.update_keyword_suggestions(e, idx))
            self.keyword_combos.append(combo)

        # Компетенции
        tk.Label(frame, text="Компетенции (до 3):").grid(row=6, column=0, sticky=tk.W, padx=10, pady=5)
        self.competencies_frame = tk.Frame(frame)
        self.competencies_frame.grid(row=6, column=1, padx=10, pady=5)

        self.competency_vars = []
        self.level_combos = []

        self.cursor.execute("SELECT id, name FROM competencies")
        comps = [f"{row[0]}: {row[1]}" for row in self.cursor.fetchall()]

        for i in range(3):
            var = tk.StringVar()
            combo = ttk.Combobox(self.competencies_frame, textvariable=var, values=comps, width=25, state="readonly")
            combo.grid(row=i, column=0, padx=5, pady=2)

            level = ttk.Combobox(self.competencies_frame, values=["1", "2", "3", "4", "5"], width=5, state="readonly")
            level.grid(row=i, column=1, padx=5, pady=2)

            self.competency_vars.append(var)
            self.level_combos.append(level)

        # Кнопка
        tk.Button(frame, text="Добавить запись", command=self.add_entry, bg="lightblue").grid(row=7, column=1, pady=20)

    def update_keyword_suggestions(self, event, index):
        current = self.keyword_combos[index].get().lower()
        if current:
            self.cursor.execute("SELECT keyword FROM keywords WHERE LOWER(keyword) LIKE %s", (f"{current}%",))
            suggestions = [row[0] for row in self.cursor.fetchall()]
            self.keyword_combos[index]['values'] = suggestions

    def add_entry(self):
        title = self.title_entry.get().strip()
        entry_type = self.type_combo.get()
        date = self.date_entry.get().strip()
        description = self.description_text.get("1.0", tk.END).strip()
        coauthors = self.coauthors_entry.get().strip()

        if not title or not entry_type or not date:
            messagebox.showerror("Ошибка", "Заполните обязательные поля")
            return

        try:
            datetime.strptime(date, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("Ошибка", "Неверный формат даты")
            return

        keywords = [combo.get().strip() for combo in self.keyword_combos if combo.get().strip()]

        competencies = []
        for i in range(3):
            comp_val = self.competency_vars[i].get()
            level_val = self.level_combos[i].get()
            if comp_val and level_val:
                comp_id = int(comp_val.split(":")[0])
                level = int(level_val)
                competencies.append((comp_id, level))

        try:
            self.cursor.execute(
                "INSERT INTO entries (title, type, date, description, coauthors, user_id) VALUES (%s, %s, %s, %s, %s, %s) RETURNING id",
                (title, entry_type, date, description, coauthors, self.current_user_id)
            )
            entry_id = self.cursor.fetchone()[0]

            for keyword in keywords:
                self.cursor.execute("INSERT INTO keywords (keyword) VALUES (%s) ON CONFLICT DO NOTHING", (keyword,))
                self.cursor.execute("SELECT id FROM keywords WHERE keyword = %s", (keyword,))
                kw_id = self.cursor.fetchone()[0]
                self.cursor.execute("INSERT INTO entry_keywords (entry_id, keyword_id) VALUES (%s, %s)",
                                    (entry_id, kw_id))

            for comp_id, level in competencies:
                self.cursor.execute(
                    "INSERT INTO entry_competencies (entry_id, competency_id, level) VALUES (%s, %s, %s)",
                    (entry_id, comp_id, level)
                )

            self.conn.commit()

            # Очистка формы
            self.title_entry.delete(0, tk.END)
            self.type_combo.set('')
            self.description_text.delete("1.0", tk.END)
            self.coauthors_entry.delete(0, tk.END)
            for combo in self.keyword_combos:
                combo.set('')
            for var in self.competency_vars:
                var.set('')
            for combo in self.level_combos:
                combo.set('')

            self.load_entries()
            self.update_statistics()
            self.check_achievements()
            messagebox.showinfo("Успех", "Запись добавлена")

        except Exception as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка", f"Ошибка БД: {str(e)}")

    def create_view_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Мои записи")

        # Панель инструментов
        toolbar = tk.Frame(frame)
        toolbar.pack(fill=tk.X, padx=5, pady=5)

        tk.Button(toolbar, text="Обновить", command=self.load_entries).pack(side=tk.LEFT, padx=5)
        tk.Button(toolbar, text="Удалить выбранное", command=self.delete_entry).pack(side=tk.LEFT, padx=5)

        # Дерево записей
        columns = ("ID", "Название", "Тип", "Дата")
        self.tree = ttk.Treeview(frame, columns=columns, show="headings", height=20)

        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100)

        scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

        self.tree.bind('<Double-1>', self.show_entry_details)

    def load_entries(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        self.cursor.execute("""
            SELECT id, title, type, date 
            FROM entries WHERE user_id = %s ORDER BY date DESC
        """, (self.current_user_id,))

        for row in self.cursor.fetchall():
            self.tree.insert("", tk.END, values=row)

    def delete_entry(self):
        selection = self.tree.selection()
        if not selection:
            return

        if messagebox.askyesno("Подтверждение", "Удалить запись?"):
            item = self.tree.item(selection[0])
            entry_id = item['values'][0]

            try:
                self.cursor.execute("DELETE FROM entries WHERE id = %s", (entry_id,))
                self.conn.commit()
                self.load_entries()
                self.update_statistics()
            except Exception as e:
                self.conn.rollback()
                messagebox.showerror("Ошибка", f"Не удалось удалить: {str(e)}")

    def show_entry_details(self, event):
        selection = self.tree.selection()
        if not selection:
            return

        item = self.tree.item(selection[0])
        entry_id = item['values'][0]

        self.cursor.execute("""
            SELECT e.title, e.type, e.date, e.description, e.coauthors,
                   COALESCE(string_agg(k.keyword, ', '), 'Нет') as keywords
            FROM entries e
            LEFT JOIN entry_keywords ek ON e.id = ek.entry_id
            LEFT JOIN keywords k ON ek.keyword_id = k.id
            WHERE e.id = %s
            GROUP BY e.id
        """, (entry_id,))

        row = self.cursor.fetchone()

        if row:
            details = f"Название: {row[0]}\n\n"
            details += f"Тип: {row[1]}\n\n"
            details += f"Дата: {row[2]}\n\n"
            details += f"Описание: {row[3]}\n\n"
            details += f"Соавторы: {row[4]}\n\n"
            details += f"Ключевые слова: {row[5]}"

            messagebox.showinfo("Детали записи", details)

    def create_research_map_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Исследовательская карта")

        # Ключевые слова
        kw_frame = tk.LabelFrame(frame, text="Ключевые слова")
        kw_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.keywords_text = tk.Text(kw_frame, height=10, state=tk.DISABLED)
        self.keywords_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Соавторы
        ca_frame = tk.LabelFrame(frame, text="Соавторы")
        ca_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.coauthors_text = tk.Text(ca_frame, height=10, state=tk.DISABLED)
        self.coauthors_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    def create_achievements_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Достижения")

        self.achievements_text = tk.Text(frame, height=25, state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(frame, command=self.achievements_text.yview)
        self.achievements_text.configure(yscrollcommand=scrollbar.set)

        self.achievements_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

    def create_competencies_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Мои компетенции")

        # Уровень компетенций
        comp_frame = tk.LabelFrame(frame, text="Уровень компетенций")
        comp_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.competencies_text = tk.Text(comp_frame, height=8, state=tk.DISABLED)
        self.competencies_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Рекомендации
        rec_frame = tk.LabelFrame(frame, text="Рекомендации")
        rec_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.recommendations_text = tk.Text(rec_frame, height=8, state=tk.DISABLED)
        self.recommendations_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    def create_goals_tab(self):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Цели на семестр")

        # Форма добавления
        form_frame = tk.Frame(frame)
        form_frame.pack(fill=tk.X, padx=10, pady=10)

        tk.Label(form_frame, text="Цель:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.goal_entry = tk.Entry(form_frame, width=40)
        self.goal_entry.grid(row=0, column=1, padx=5, pady=5)

        tk.Label(form_frame, text="Целевое значение:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.target_entry = tk.Entry(form_frame, width=10)
        self.target_entry.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)

        tk.Button(form_frame, text="Добавить цель", command=self.add_goal).grid(row=2, column=1, pady=10)

        # Список целей
        goals_frame = tk.LabelFrame(frame, text="Текущие цели")
        goals_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.goals_text = tk.Text(goals_frame, height=12, state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(goals_frame, command=self.goals_text.yview)
        self.goals_text.configure(yscrollcommand=scrollbar.set)

        self.goals_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

    def add_goal(self):
        desc = self.goal_entry.get().strip()
        target = self.target_entry.get().strip()

        if not desc:
            messagebox.showerror("Ошибка", "Введите описание цели")
            return

        try:
            target_val = int(target) if target.isdigit() else 1
        except:
            target_val = 1

        try:
            self.cursor.execute(
                "INSERT INTO goals (description, target_value, current_value, user_id) VALUES (%s, %s, %s, %s)",
                (desc, target_val, 0, self.current_user_id)
            )
            self.conn.commit()

            self.goal_entry.delete(0, tk.END)
            self.target_entry.delete(0, tk.END)
            self.load_goals()

        except Exception as e:
            self.conn.rollback()
            messagebox.showerror("Ошибка", f"Не удалось добавить цель: {str(e)}")

    def load_goals(self):
        self.cursor.execute(
            "SELECT description, target_value, current_value FROM goals WHERE user_id = %s",
            (self.current_user_id,)
        )

        content = ""
        for desc, target, current in self.cursor.fetchall():
            status = "✅ Выполнено" if current >= target else "🔄 В процессе"
            content += f"{desc}\nПрогресс: {current}/{target} - {status}\n\n"

        self.goals_text.config(state=tk.NORMAL)
        self.goals_text.delete("1.0", tk.END)
        self.goals_text.insert("1.0", content if content else "Цели не установлены")
        self.goals_text.config(state=tk.DISABLED)

    def update_statistics(self):
        # Ключевые слова
        self.cursor.execute("""
            SELECT k.keyword, COUNT(ek.entry_id) as count
            FROM keywords k
            JOIN entry_keywords ek ON k.id = ek.keyword_id
            JOIN entries e ON ek.entry_id = e.id
            WHERE e.user_id = %s
            GROUP BY k.keyword
            ORDER BY count DESC
        """, (self.current_user_id,))

        kw_content = ""
        for kw, count in self.cursor.fetchall():
            kw_content += f"{kw} — {count} записей\n"

        self.keywords_text.config(state=tk.NORMAL)
        self.keywords_text.delete("1.0", tk.END)
        self.keywords_text.insert("1.0", kw_content if kw_content else "Нет данных")
        self.keywords_text.config(state=tk.DISABLED)

        # Соавторы
        self.cursor.execute("""
            SELECT e.coauthors FROM entries e
            WHERE e.user_id = %s AND e.coauthors IS NOT NULL AND e.coauthors != ''
        """, (self.current_user_id,))

        coauthors_dict = {}
        for row in self.cursor.fetchall():
            for ca in row[0].split(","):
                ca = ca.strip()
                if ca:
                    coauthors_dict[ca] = coauthors_dict.get(ca, 0) + 1

        ca_content = ""
        for ca, count in sorted(coauthors_dict.items(), key=lambda x: x[1], reverse=True):
            ca_content += f"{ca} — {count} работ\n"

        self.coauthors_text.config(state=tk.NORMAL)
        self.coauthors_text.delete("1.0", tk.END)
        self.coauthors_text.insert("1.0", ca_content if ca_content else "Нет данных")
        self.coauthors_text.config(state=tk.DISABLED)

        # Компетенции
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL
            ORDER BY avg_level DESC
        """, (self.current_user_id,))

        comp_content = ""
        weak_content = ""
        rec_content = ""

        for name, level in self.cursor.fetchall():
            level = float(level)
            comp_content += f"{name}: {level:.2f}\n"

            if level < 3:
                weak_content += f"{name}: {level:.2f}\n"

                if "Презентация" in name:
                    rec_content += "Рекомендуется выступить на студенческой конференции\n"
                elif "Командная" in name:
                    rec_content += "Рекомендуется участвовать в групповых проектах\n"
                elif "БД" in name:
                    rec_content += "Рекомендуется пройти курс по базам данных\n"

        if not rec_content:
            rec_content = "Все компетенции развиты хорошо"

        self.competencies_text.config(state=tk.NORMAL)
        self.competencies_text.delete("1.0", tk.END)
        self.competencies_text.insert("1.0", comp_content if comp_content else "Нет данных")
        self.competencies_text.config(state=tk.DISABLED)

        self.recommendations_text.config(state=tk.NORMAL)
        self.recommendations_text.delete("1.0", tk.END)
        self.recommendations_text.insert("1.0", rec_content)
        self.recommendations_text.config(state=tk.DISABLED)

        # Достижения
        self.cursor.execute(
            "SELECT name, description, unlocked_date FROM achievements WHERE user_id = %s",
            (self.current_user_id,)
        )

        ach_content = ""
        for name, desc, date in self.cursor.fetchall():
            ach_content += f"{name}\n{desc}\nПолучено: {date}\n\n"

        self.achievements_text.config(state=tk.NORMAL)
        self.achievements_text.delete("1.0", tk.END)
        self.achievements_text.insert("1.0", ach_content if ach_content else "Достижений нет")
        self.achievements_text.config(state=tk.DISABLED)

        # Цели
        self.load_goals()

    def check_achievements(self):
        self.cursor.execute("SELECT COUNT(*) FROM entries WHERE user_id = %s", (self.current_user_id,))
        total = self.cursor.fetchone()[0]

        if total == 1:
            self.unlock_achievement("Первый шаг", "Создана первая запись")

        self.cursor.execute("""
            SELECT COUNT(*) FROM entries 
            WHERE user_id = %s AND coauthors IS NOT NULL AND coauthors != ''
        """, (self.current_user_id,))

        with_coauthors = self.cursor.fetchone()[0]
        if with_coauthors >= 3:
            self.unlock_achievement("Командный игрок", "Три и более записи с соавторами")

        self.cursor.execute("SELECT COUNT(DISTINCT type) FROM entries WHERE user_id = %s", (self.current_user_id,))
        types = self.cursor.fetchone()[0]
        if types >= 3:
            self.unlock_achievement("Разносторонний", "Записи минимум трёх разных типов")

        self.cursor.execute("""
            SELECT EXTRACT(YEAR FROM date) as year, COUNT(*) as count
            FROM entries WHERE user_id = %s
            GROUP BY EXTRACT(YEAR FROM date)
            HAVING COUNT(*) >= 3
        """, (self.current_user_id,))

        if self.cursor.fetchone():
            self.unlock_achievement("Плодотворный год", "Три и более записи за один календарный год")

        self.cursor.execute("SELECT SUM(LENGTH(description)) FROM entries WHERE user_id = %s", (self.current_user_id,))
        chars = self.cursor.fetchone()[0] or 0
        if chars > 5000:
            self.unlock_achievement("Словобог", "Суммарный объём описаний превысил 5000 символов")

    def unlock_achievement(self, name, desc):
        self.cursor.execute(
            "SELECT id FROM achievements WHERE name = %s AND user_id = %s",
            (name, self.current_user_id)
        )

        if not self.cursor.fetchone():
            self.cursor.execute(
                "INSERT INTO achievements (name, description, user_id, unlocked_date) VALUES (%s, %s, %s, %s)",
                (name, desc, self.current_user_id, datetime.now().strftime("%Y-%m-%d"))
            )
            self.conn.commit()
            self.update_statistics()

    def export_to_word(self):
        doc = Document()

        # Заголовок
        title = doc.add_paragraph()
        title_run = title.add_run('Отчёт по портфолио')
        title_run.font.size = Pt(16)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Пользователь ID: {self.current_user_id}')

        # Записи
        doc.add_heading('Записи портфолио', level=1)
        self.cursor.execute(
            "SELECT title, type, date, description, coauthors FROM entries WHERE user_id = %s ORDER BY date DESC",
            (self.current_user_id,)
        )

        for title, type_, date, desc, coauthors in self.cursor.fetchall():
            doc.add_heading(title, level=2)
            doc.add_paragraph(f'Тип: {type_}')
            doc.add_paragraph(f'Дата: {date}')
            doc.add_paragraph(f'Соавторы: {coauthors if coauthors else "Нет"}')
            doc.add_paragraph(f'Описание: {desc if desc else "Нет"}')

        # Ключевые слова
        doc.add_heading('Ключевые слова', level=1)
        self.cursor.execute("""
            SELECT k.keyword, COUNT(ek.entry_id) as count
            FROM keywords k
            JOIN entry_keywords ek ON k.id = ek.keyword_id
            JOIN entries e ON ek.entry_id = e.id
            WHERE e.user_id = %s
            GROUP BY k.keyword
            ORDER BY count DESC
        """, (self.current_user_id,))

        for kw, count in self.cursor.fetchall():
            doc.add_paragraph(f'{kw} — {count} записей')

        # Соавторы
        doc.add_heading('Соавторы', level=1)
        self.cursor.execute("""
            SELECT e.coauthors FROM entries e
            WHERE e.user_id = %s AND e.coauthors IS NOT NULL AND e.coauthors != ''
        """, (self.current_user_id,))

        coauthors_dict = {}
        for row in self.cursor.fetchall():
            for ca in row[0].split(","):
                ca = ca.strip()
                if ca:
                    coauthors_dict[ca] = coauthors_dict.get(ca, 0) + 1

        for ca, count in sorted(coauthors_dict.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f'{ca} — {count} работ')

        # Компетенции
        doc.add_heading('Компетенции', level=1)
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL
            ORDER BY avg_level DESC
        """, (self.current_user_id,))

        for name, level in self.cursor.fetchall():
            doc.add_paragraph(f'{name}: {float(level):.2f}')

        # Рекомендации
        doc.add_heading('Рекомендации', level=1)
        self.cursor.execute("""
            SELECT c.name, CAST(AVG(ec.level) AS DECIMAL(10,2)) as avg_level
            FROM competencies c
            LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
            LEFT JOIN entries e ON ec.entry_id = e.id AND e.user_id = %s
            GROUP BY c.id, c.name
            HAVING AVG(ec.level) IS NOT NULL AND AVG(ec.level) < 3
        """, (self.current_user_id,))

        weak = self.cursor.fetchall()
        if weak:
            doc.add_paragraph('Рекомендации по развитию слабых зон:')
            for name, level in weak:
                level = float(level)
                if "Презентация" in name:
                    doc.add_paragraph(f'- {name} ({level:.2f}): выступите на студенческой конференции')
                elif "Командная" in name:
                    doc.add_paragraph(f'- {name} ({level:.2f}): участвуйте в групповых проектах')
                elif "БД" in name:
                    doc.add_paragraph(f'- {name} ({level:.2f}): пройдите курс по базам данных')
        else:
            doc.add_paragraph('Все компетенции развиты хорошо')

        # Достижения
        doc.add_heading('Достижения', level=1)
        self.cursor.execute(
            "SELECT name, description, unlocked_date FROM achievements WHERE user_id = %s",
            (self.current_user_id,)
        )

        for name, desc, date in self.cursor.fetchall():
            doc.add_paragraph(f'● {name}: {desc} ({date})')

        # Сохранение
        filename = f"portfolio_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(filename)
        messagebox.showinfo("Экспорт завершён", f"Отчёт сохранён: {filename}")

    def __del__(self):
        if hasattr(self, 'cursor'):
            self.cursor.close()
        if hasattr(self, 'conn'):
            self.conn.close()


if __name__ == "__main__":
    root = tk.Tk()
    app = PortfolioApp(root)
    root.mainloop()